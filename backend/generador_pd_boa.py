import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert

def _add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h

def _set_cell_background(cell, color_hex):
    # Set background color of a cell using oxml
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def generate(data, out_docx, out_pdf):
    doc = Document()
    
    # --- Portada ---
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROGRAMACIÓN DIDÁCTICA")
    run.font.size = Pt(28)
    run.font.bold = True
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run(data.get("modulo", "Módulo Profesional"))
    run2.font.size = Pt(20)
    
    doc.add_paragraph()
    
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.add_run(f"Ciclo: {data.get('ciclo', '')}\n")
    info_p.add_run(f"Departamento: {data.get('departamento', '')}\n")
    info_p.add_run(f"Curso Académico: {data.get('curso_academico', '')}\n")
    info_p.add_run(f"Horas Totales: {data.get('horas_totales', '')} h\n")
    
    doc.add_page_break()
    
    # --- 1. Resultados de Aprendizaje y Criterios de Evaluación ---
    _add_heading(doc, "1. Resultados de Aprendizaje y Criterios de Evaluación", 1)
    df_ra = data.get("df_ra", [])
    if df_ra:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'RA'
        hdr_cells[1].text = 'Descripción'
        hdr_cells[2].text = 'Peso %'
        _set_cell_background(hdr_cells[0], 'DDDDDD')
        _set_cell_background(hdr_cells[1], 'DDDDDD')
        _set_cell_background(hdr_cells[2], 'DDDDDD')
        
        for ra in df_ra:
            row_cells = table.add_row().cells
            row_cells[0].text = str(ra.get("id_ra", ""))
            row_cells[1].text = str(ra.get("desc_ra", ""))
            row_cells[2].text = f"{ra.get('peso_ra', 0)}%"
    else:
        doc.add_paragraph("No hay resultados de aprendizaje definidos.")
        
    doc.add_page_break()
    
    # --- 2. Secuenciación y Unidades Didácticas ---
    _add_heading(doc, "2. Secuenciación y Unidades Didácticas", 1)
    df_ud = data.get("df_ud", [])
    if df_ud:
        for ud in df_ud:
            _add_heading(doc, f"Unidad {ud.get('id_ud', '')}: {ud.get('desc_ud', '')}", 2)
            p = doc.add_paragraph()
            p.add_run("Horas asignadas: ").bold = True
            p.add_run(f"{ud.get('horas_ud', 0)}h\n")
            
            p.add_run("Intención Educativa: ").bold = True
            p.add_run(f"{ud.get('Intencion_Educativa', 'No especificada')}\n")
            
            p.add_run("Temporización: ").bold = True
            p.add_run(f"{ud.get('Temporizacion', 'No especificada')}\n")
            
            p.add_run("Agrupamientos: ").bold = True
            p.add_run(f"{ud.get('Agrupamientos', 'No especificados')}\n")
    else:
        doc.add_paragraph("No hay unidades didácticas definidas.")
        
    doc.add_page_break()
    
    # --- 3. Tareas Competenciales ---
    _add_heading(doc, "3. Tareas Competenciales", 1)
    df_act = data.get("df_act", [])
    if df_act:
        for act in df_act:
            _add_heading(doc, f"Tarea: {act.get('Nombre_Tarea', act.get('ID', ''))}", 2)
            p = doc.add_paragraph()
            p.add_run("Briefing / Contexto Profesional: ").bold = True
            p.add_run(f"\n{act.get('Briefing', 'No especificado')}\n")
            
            p.add_run("Pasos / Desarrollo: ").bold = True
            p.add_run(f"\n{act.get('Pasos', 'No especificados')}\n")
            
            p.add_run("Evidencias: ").bold = True
            p.add_run(f"\n{act.get('Evidencias', 'No especificadas')}\n")
    else:
        doc.add_paragraph("No hay tareas competenciales definidas.")
        
    doc.add_page_break()
    
    # --- 4. Criterios de Calificación ---
    _add_heading(doc, "4. Criterios de Calificación", 1)
    redondeo = data.get("config_redondeo", {})
    doc.add_paragraph(f"Nota para aprobar: {redondeo.get('nota_aprobado', 5.0)}")
    doc.add_paragraph(f"Umbral de redondeo al alza: {redondeo.get('umbral_redondeo', 4.5)}")

    # Save DOCX
    doc.save(out_docx)
    
    # Convert to PDF
    try:
        # On Windows, docx2pdf uses COM and requires absolute paths
        abs_docx = os.path.abspath(out_docx)
        abs_pdf = os.path.abspath(out_pdf)
        convert(abs_docx, abs_pdf)
    except Exception as e:
        print(f"Error converting to PDF: {e}")
        # If conversion fails, we'll just skip the PDF (handled in routers/pdf.py by returning DOCX instead)
        if os.path.exists(out_pdf):
            os.remove(out_pdf)
