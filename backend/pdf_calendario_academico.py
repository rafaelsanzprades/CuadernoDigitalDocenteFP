import io
import os
import calendar
from datetime import date
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import BaseDocTemplate, Frame, PageTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas as pdfcanvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

NOMBRE_MESES = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
                "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]

# Arial de verdad si el sistema la tiene (Windows), si no cae a Helvetica
# (equivalente base14 de PDF, siempre disponible, incluido en Cloud Run).
try:
    _WIN_FONTS = "C:/Windows/Fonts"
    pdfmetrics.registerFont(TTFont("Arial", os.path.join(_WIN_FONTS, "arial.ttf")))
    pdfmetrics.registerFont(TTFont("Arial-Bold", os.path.join(_WIN_FONTS, "arialbd.ttf")))
    pdfmetrics.registerFont(TTFont("Arial-Italic", os.path.join(_WIN_FONTS, "ariali.ttf")))
    pdfmetrics.registerFont(TTFont("Arial-BoldItalic", os.path.join(_WIN_FONTS, "arialbi.ttf")))
    pdfmetrics.registerFontFamily("Arial", normal="Arial", bold="Arial-Bold",
                                   italic="Arial-Italic", boldItalic="Arial-BoldItalic")
    FONT_R, FONT_B, FONT_I, FONT_BI = "Arial", "Arial-Bold", "Arial-Italic", "Arial-BoldItalic"
except Exception:
    FONT_R, FONT_B, FONT_I, FONT_BI = "Helvetica", "Helvetica-Bold", "Helvetica-Oblique", "Helvetica-BoldOblique"

# ------------------------------------------------------------------ #
#  Función que se llama en CADA página para dibujar cabecera y pie   #
# ------------------------------------------------------------------ #
def _draw_page_decorations(canv, doc):
    canv.saveState()
    W, H = landscape(A4)

    # ---- CABECERA: Título centrado ----
    canv.setFont(FONT_B, 10)
    canv.setFillColor(colors.HexColor("#777777"))
    canv.drawCentredString(W / 2, H - 1.5 * cm, doc.cal_titulo)

    # ---- PIE: Referencia abajo a la derecha ----
    canv.setFont(FONT_R, 9)
    canv.setFillColor(colors.HexColor("#777777"))
    canv.drawRightString(W - 1 * cm, 1 * cm, doc.cal_pie)

    canv.restoreState()

# ------------------------------------------------------------------ #
#  Función principal                                                  #
# ------------------------------------------------------------------ #
def generar_pdf_calendario(info_modulo, info_fechas, planning_ledger, calendar_notes):
    buffer = io.BytesIO()

    W, H = landscape(A4)
    margin = 2 * cm
    # Espacio extra arriba (cabecera) y abajo (pie) para no solapar
    top_margin    = 2 * cm
    bottom_margin = 2 * cm

    doc = BaseDocTemplate(
        buffer,
        pagesize=landscape(A4),
        rightMargin=margin, leftMargin=margin,
        topMargin=top_margin, bottomMargin=bottom_margin,
    )

    # Guardamos los textos de cabecera/pie como atributos del doc para
    # que la función de decoración pueda acceder a ellos
    ini = info_fechas.get("ini_1t", date(2025, 9, 1))
    fin = info_fechas.get("fin_3t", date(2026, 6, 30))
    ini_feoe = info_fechas.get("ini_feoe", date(2026, 3, 16))
    fin_feoe = info_fechas.get("fin_feoe", date(2026, 5, 29))
    
    doc.cal_titulo = f"Calendario Académico {ini.year} - {fin.year}. {info_modulo.get('modulo', 'Módulo')}"
    doc.cal_pie    = f"{info_modulo.get('centro', 'IES Andalán')} ({info_modulo.get('profesorado', 'Rafael Sanz Prades')})"

    frame = Frame(margin, bottom_margin, W - 2*margin, H - top_margin - bottom_margin, id='main')
    template = PageTemplate(id='cal', frames=[frame], onPage=_draw_page_decorations)
    doc.addPageTemplates([template])

    elements = []

    # ---- Meses del curso ----
    start_date = ini.replace(day=1)
    meses_curso = []
    curr = start_date
    while curr <= fin:
        meses_curso.append((curr.year, curr.month))
        if curr.month == 12:
            curr = curr.replace(year=curr.year + 1, month=1, day=1)
        else:
            curr = curr.replace(month=curr.month + 1, day=1)

    # ---- Funciones auxiliares ----
    def get_cell_data(year, month, day):
        if day == 0:
            return ("", "", "", "", False, "")
        
        d_str = f"{day:02d}/{month:02d}/{year}"
        fecha_obj = date(year, month, day)
        es_finde = fecha_obj.weekday() >= 5
        
        # Festivos
        desc_festivo = calendar_notes.get(f"f_{d_str}", "").strip()
        es_festivo = es_finde or bool(desc_festivo)
        
        # UD
        uds = planning_ledger.get(d_str, [])
        texto_ud = ", ".join(uds)[:20] if uds else ""
        
        # Relevantes
        desc_rel = calendar_notes.get(f"r_{d_str}", "").strip()[:30]
        
        # FEOE: auto-calculado basado en el rango global, sin fines de semana
        desc_feoe = "FEOE" if (ini_feoe <= fecha_obj <= fin_feoe and not es_finde) else ""
        
        texto_plan_ud = texto_ud
            
        texto_plan_rel = desc_rel
        
        return (f"{day:02d}", texto_plan_ud[:30], texto_plan_rel, desc_festivo[:30], es_festivo, desc_feoe)

    FESTIVO_BG = colors.HexColor("#fdecea")
    UD_BG = colors.HexColor("#ede7f6")

    def build_day_cell(day_str, tfeoe):
        """Fila 0 de la casilla: solo el número de día, siempre centrado
        (fondo rojizo aparte, vía festivo_bg_cells, si es festivo)."""
        if not day_str:
            return ""
        texto = f"<b>{day_str}</b>"
        if tfeoe:
            texto += f" <font size=8>{tfeoe}</font>"
        style = ParagraphStyle(name='D', alignment=1, fontSize=14, fontName=FONT_B, textColor=colors.black)
        return Paragraph(texto, style)

    def get_month_grid(year, month):
        cal = calendar.monthcalendar(year, month)
        g_dias, g_festivo, g_ud, g_rel, g_fest = [], [], [], [], []
        for week in cal:
            valid_day = next(d for d in week if d != 0)
            week_num  = date(year, month, valid_day).isocalendar()[1]

            fila_dias    = [str(week_num)]
            fila_festivo = [""]
            fila_ud      = [""]
            fila_rel     = [""]
            fila_fest    = [False]

            for col_idx, d in enumerate(week):
                td, tud, trel, tfest_desc, ef, tfeoe = get_cell_data(year, month, d)
                # 4 filas por casilla: día (con fondo rojizo si es festivo) /
                # festivo (texto, propio, centrado) / relevante (a la
                # derecha) / docencia-UD (a la izquierda, fondo malva).
                fila_dias.append(build_day_cell(td, tfeoe) if td else "")
                fila_festivo.append(tfest_desc if ef else "")
                fila_ud.append("" if ef else tud)
                fila_rel.append("" if ef else trel)
                fila_fest.append(ef)

            g_dias.append(fila_dias)
            g_festivo.append(fila_festivo)
            g_ud.append(fila_ud)
            g_rel.append(fila_rel)
            g_fest.append(fila_fest)

        return g_dias, g_festivo, g_ud, g_rel, g_fest

    # ---- Alturas fijas de fila ----
    ROW_MES     = 2.5 * cm  # Fila del nombre del mes  (≈ 3× la normal)
    ROW_HEAD    = 0.7 * cm  # Fila de Lun/Mar/Mié…
    ROW_DIAS    = 0.6 * cm  # Fila de número de día
    ROW_FESTIVO = 0.5 * cm  # Fila de nombre del festivo
    ROW_REL     = 0.5 * cm  # Fila de fechas relevantes
    ROW_UD      = 0.5 * cm  # Fila de docencia (UD)

    # ---- Anchos de columnas ----
    colWidths = [1.5 * cm] + [4.75 * cm] * 5 + [1.5 * cm] * 2

    # ---- Construir tabla por mes ----
    for mes_idx, (year, month) in enumerate(meses_curso):
        g_dias, g_festivo, g_ud, g_rel, g_fest = get_month_grid(year, month)
        num_weeks = len(g_dias)

        t_data = []

        # Fila 0 - nombre del mes
        t_data.append([f"{NOMBRE_MESES[month-1]}  {year}", "", "", "", "", "", "", ""])

        # Fila 1 - cabecera días
        t_data.append(["Sem.", "Lunes", "Martes", "Miércoles", "Jueves", "Viernes", "Sáb.", "Dom."])

        # Filas de datos
        festivo_bg_cells = []
        ud_bg_cells = []
        row_counter = 2

        for w in range(num_weeks):
            t_data.append(g_dias[w])
            t_data.append(g_festivo[w])
            t_data.append(g_rel[w])
            t_data.append(g_ud[w])

            for c in range(1, 8):
                if g_fest[w][c]:
                    festivo_bg_cells.append((c, row_counter))
                    festivo_bg_cells.append((c, row_counter + 1))
                elif g_dias[w][c] != "" and g_ud[w][c]:
                    # Fondo malva solo si ese día concreto tiene UD asignada.
                    ud_bg_cells.append((c, row_counter + 3))

            row_counter += 4

        # Alturas: [mes, head, (dias, festivo, rel, ud) × num_weeks]
        row_heights = [ROW_MES, ROW_HEAD]
        for _ in range(num_weeks):
            row_heights.extend([ROW_DIAS, ROW_FESTIVO, ROW_REL, ROW_UD])

        t = Table(t_data, colWidths=colWidths, rowHeights=row_heights)

        style_list = [
            # ---- Fila del mes ----
            ('SPAN',       (0, 0), (7, 0)),
            ('ALIGN',      (0, 0), (7, 0), 'CENTER'),
            ('VALIGN',     (0, 0), (7, 0), 'MIDDLE'),
            ('FONTNAME',   (0, 0), (7, 0), FONT_B),
            ('FONTSIZE',   (0, 0), (7, 0), 22),          # 2× cabecera
            ('BACKGROUND', (0, 0), (7, 0), colors.white),
            ('TEXTCOLOR',  (0, 0), (7, 0), colors.black),

            # ---- Fila cabecera días ----
            ('ALIGN',      (0, 1), (-1, 1), 'CENTER'),
            ('VALIGN',     (0, 1), (-1, 1), 'MIDDLE'),
            ('FONTNAME',   (0, 1), (-1, 1), FONT_B),
            ('FONTSIZE',   (0, 1), (-1, 1), 10),
            ('BACKGROUND', (0, 1), (7, 1), colors.HexColor("#e0e0e0")),
            ('TEXTCOLOR',  (0, 1), (-1, 1), colors.black),

            # ---- Cuadrícula general ----
            ('BOX',       (0, 0), (-1, -1), 0.5, colors.HexColor("#bbbbbb")),
            ('LINEAFTER', (0, 0), (-1, -1), 0.5, colors.HexColor("#bbbbbb")),
            ('LINEBELOW', (0, 0), (-1, 1), 0.5, colors.HexColor("#bbbbbb")),

            # ---- Columna "Sem" ----
            ('BACKGROUND', (0, 2), (0, -1), colors.HexColor("#f5f5f5")),
            ('TEXTCOLOR',  (0, 2), (0, -1), colors.black),
            ('ALIGN',      (0, 2), (0, -1), 'CENTER'),
            ('VALIGN',     (0, 2), (0, -1), 'MIDDLE'),
            ('FONTNAME',   (0, 2), (0, -1), FONT_B),
            ('FONTSIZE',   (0, 2), (0, -1), 16),
        ]

        # ---- Estilos por fila de datos ----
        for r_idx in range(2, len(t_data), 4):
            # Fila 0: número de día
            style_list += [
                ('ALIGN',    (1, r_idx), (-1, r_idx), 'CENTER'),
                ('VALIGN',   (1, r_idx), (-1, r_idx), 'MIDDLE'),
                # Fusionar columna Sem verticalmente con las 3 filas debajo
                ('SPAN',     (0, r_idx), (0, r_idx + 3)),
            ]

            # Fila 1: festivo, centrado
            style_list += [
                ('FONTNAME',  (1, r_idx+1), (-1, r_idx+1), FONT_B),
                ('FONTSIZE',  (1, r_idx+1), (-1, r_idx+1), 8),
                ('TEXTCOLOR', (1, r_idx+1), (-1, r_idx+1), colors.black),
                ('ALIGN',     (1, r_idx+1), (-1, r_idx+1), 'CENTER'),
                ('VALIGN',    (1, r_idx+1), (-1, r_idx+1), 'MIDDLE'),
            ]

            # Fila 2: relevante, alineado a la derecha
            style_list += [
                ('FONTNAME',  (1, r_idx+2), (-1, r_idx+2), FONT_I),
                ('FONTSIZE',  (1, r_idx+2), (-1, r_idx+2), 8),
                ('TEXTCOLOR', (1, r_idx+2), (-1, r_idx+2), colors.black),
                ('ALIGN',     (1, r_idx+2), (-1, r_idx+2), 'RIGHT'),
                ('VALIGN',    (1, r_idx+2), (-1, r_idx+2), 'MIDDLE'),
            ]

            # Fila 3: docencia (UD), alineado a la izquierda, fondo malva
            style_list += [
                ('FONTNAME',  (1, r_idx+3), (-1, r_idx+3), FONT_R),
                ('FONTSIZE',  (1, r_idx+3), (-1, r_idx+3), 9),
                ('TEXTCOLOR', (1, r_idx+3), (-1, r_idx+3), colors.black),
                ('ALIGN',     (1, r_idx+3), (-1, r_idx+3), 'LEFT'),
                ('VALIGN',    (1, r_idx+3), (-1, r_idx+3), 'MIDDLE'),
                # Línea divisoria de semana separando bloques diarios
                ('LINEBELOW', (0, r_idx+3), (-1, r_idx+3), 0.5, colors.HexColor("#bbbbbb")),
            ]

        # ---- Fondo rojizo (festivo, filas 0-1) y malva (docencia, fila 3) ----
        for (col, row) in festivo_bg_cells:
            style_list.append(('BACKGROUND', (col, row), (col, row), FESTIVO_BG))
        for (col, row) in ud_bg_cells:
            style_list.append(('BACKGROUND', (col, row), (col, row), UD_BG))

        t.setStyle(TableStyle(style_list))
        elements.append(t)

        if mes_idx < len(meses_curso) - 1:
            elements.append(PageBreak())

    doc.build(elements)
    buffer.seek(0)
    return buffer


DIAS_SEMANA = ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes", "Sábado", "Domingo"]

_CAL_HEADERS = ["Sem.", "Lunes", "Martes", "Miércoles", "Jueves", "Viernes", "Sáb.", "Dom."]
_CAL_FESTIVO_BG = "FDECEA"
_CAL_HEADER_BG = "E0E0E0"
_CAL_SEM_BG = "F5F5F5"
_CAL_UD_BG = "EDE7F6"


def generar_docx_calendario(info_modulo, info_fechas, planning_ledger, calendar_notes):
    """Versión .docx editable de la misma cuadrícula visual que el PDF
    (una tabla por mes, columnas Sem./Lun-Dom, festivos resaltados) en vez
    de una lista plana de filas, para que ambos formatos se vean igual."""
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx_helpers import new_document, add_title, add_meta_line, doc_to_bytes

    ini = info_fechas.get("ini_1t", date(2025, 9, 1))
    fin = info_fechas.get("fin_3t", date(2026, 6, 30))
    ini_feoe = info_fechas.get("ini_feoe", date(2026, 3, 16))
    fin_feoe = info_fechas.get("fin_feoe", date(2026, 5, 29))

    doc = new_document(landscape=True)
    add_title(doc, f"Calendario Académico {ini.year} - {fin.year}", info_modulo.get("modulo", "Módulo"))
    add_meta_line(doc, f"{info_modulo.get('centro', '')} ({info_modulo.get('profesorado', '')})")

    col_widths = [Cm(1.6), Cm(3.6), Cm(3.6), Cm(3.6), Cm(3.6), Cm(3.6), Cm(1.6), Cm(1.6)]

    def shade(cell, hex_color):
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shd)

    def set_cell(cell, text, bold=False, italic=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = "Arial"
        return p

    def get_day_info(year, month, day):
        if day == 0:
            return None
        d_str = f"{day:02d}/{month:02d}/{year}"
        fecha_obj = date(year, month, day)
        es_finde = fecha_obj.weekday() >= 5
        desc_festivo = calendar_notes.get(f"f_{d_str}", "").strip()
        uds = planning_ledger.get(d_str, [])
        return {
            "day": day,
            "es_festivo": es_finde or bool(desc_festivo),
            "desc_festivo": desc_festivo,
            "ud": ", ".join(uds),
            "rel": calendar_notes.get(f"r_{d_str}", "").strip(),
            "feoe": "FEOE" if (ini_feoe <= fecha_obj <= fin_feoe and not es_finde) else "",
        }

    curr = ini.replace(day=1)
    first_month = True
    while curr <= fin:
        year, month = curr.year, curr.month
        if not first_month:
            doc.add_page_break()
        first_month = False

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{NOMBRE_MESES[month - 1]}  {year}")
        run.bold = True
        run.font.size = Pt(20)
        run.font.name = "Arial"

        semanas = calendar.monthcalendar(year, month)
        table = doc.add_table(rows=1 + 4 * len(semanas), cols=8)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = False

        for c, h in enumerate(_CAL_HEADERS):
            cell = table.rows[0].cells[c]
            set_cell(cell, h, bold=True, size=9)
            shade(cell, _CAL_HEADER_BG)

        for w, week in enumerate(semanas):
            base_row = 1 + w * 4
            valid_day = next(d for d in week if d != 0)
            week_num = date(year, month, valid_day).isocalendar()[1]

            sem_cell = table.rows[base_row].cells[0]
            for r in range(1, 4):
                sem_cell = sem_cell.merge(table.rows[base_row + r].cells[0])
            set_cell(sem_cell, str(week_num), bold=True, size=14)
            shade(sem_cell, _CAL_SEM_BG)

            for col_idx, day in enumerate(week):
                c = col_idx + 1
                info = get_day_info(year, month, day)
                if info is None:
                    continue
                # Fila 0: día (fondo rojizo si es festivo). Fila 1: nombre
                # del festivo, centrado. Fila 2: relevante, a la derecha.
                # Fila 3: docencia (UD), a la izquierda, fondo malva.
                cell_dia = table.rows[base_row].cells[c]
                cell_festivo = table.rows[base_row + 1].cells[c]
                cell_rel = table.rows[base_row + 2].cells[c]
                cell_ud = table.rows[base_row + 3].cells[c]

                if info["es_festivo"]:
                    set_cell(cell_dia, f"{info['day']:02d}", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
                    shade(cell_dia, _CAL_FESTIVO_BG)
                    if info["desc_festivo"]:
                        set_cell(cell_festivo, info["desc_festivo"], bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                    shade(cell_festivo, _CAL_FESTIVO_BG)
                else:
                    dia_p = set_cell(cell_dia, f"{info['day']:02d}", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
                    if info["feoe"]:
                        feoe_run = dia_p.add_run(f"  {info['feoe']}")
                        feoe_run.font.size = Pt(8)
                        feoe_run.font.name = "Arial"
                    if info["rel"]:
                        set_cell(cell_rel, info["rel"], italic=True, size=7, align=WD_ALIGN_PARAGRAPH.RIGHT)
                    set_cell(cell_ud, info["ud"], size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
                    if info["ud"]:
                        shade(cell_ud, _CAL_UD_BG)

        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w

        if curr.month == 12:
            curr = curr.replace(year=curr.year + 1, month=1, day=1)
        else:
            curr = curr.replace(month=curr.month + 1, day=1)

    return doc_to_bytes(doc)

