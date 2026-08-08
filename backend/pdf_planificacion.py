# -*- coding: utf-8 -*-
"""
pdf_planificacion.py
PDF A4 apaisado - Planificación. Horas previstas frente a impartidas.
"""
import io
import math
import pandas as pd
from datetime import datetime, date, timedelta
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import (
    BaseDocTemplate, Frame, PageTemplate,
    Table, TableStyle, Paragraph, Spacer
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

def _draw_page_decorations(canv, doc):
    canv.saveState()
    W, H = landscape(A4)
    canv.setFont("Helvetica-Bold", 10)
    canv.setFillColor(colors.HexColor("#777777"))
    canv.drawCentredString(W / 2, H - 1.5 * cm, doc.cal_titulo)
    canv.setFont("Helvetica", 9)
    canv.drawRightString(W - 1 * cm, 1 * cm, doc.cal_pie)
    canv.restoreState()


def generar_pdf_planificacion(
    info_modulo: dict,
    df_ud: pd.DataFrame,
    df_sgmt: pd.DataFrame,
    daily_ledger: dict,
    horario: dict,
    info_fechas: dict,
    calendar_notes: dict
):
    buffer = io.BytesIO()
    W, H = landscape(A4)
    left_m   = 2.0 * cm
    right_m  = 1.0 * cm
    top_m    = 2.0 * cm
    bottom_m = 1.5 * cm

    doc = BaseDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=left_m, rightMargin=right_m,
        topMargin=top_m, bottomMargin=bottom_m,
    )

    nombre_modulo = info_modulo.get("modulo", "Módulo")
    doc.cal_titulo = f"Planificación: Horas previstas frente a impartidas  ·  {nombre_modulo}"
    doc.cal_pie    = f"{info_modulo.get('centro', '')} ({info_modulo.get('profesorado', '')})"

    frame = Frame(left_m, bottom_m, W - left_m - right_m, H - top_m - bottom_m, id="main")
    doc.addPageTemplates([PageTemplate(id="port", frames=[frame], onPage=_draw_page_decorations)])

    styles = getSampleStyleSheet()
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14, spaceAfter=8, textColor=colors.black, spaceBefore=4)
    norm = ParagraphStyle("Nor", parent=styles["Normal"], fontSize=9, leading=10, alignment=TA_CENTER)
    normB = ParagraphStyle("NorB", parent=styles["Normal"], fontSize=9, leading=10, fontName="Helvetica-Bold", alignment=TA_CENTER)
    sml = ParagraphStyle("Sm", parent=styles["Normal"], fontSize=8, leading=9, alignment=TA_CENTER)
    smlB = ParagraphStyle("SmB", parent=styles["Normal"], fontSize=8, leading=9, fontName="Helvetica-Bold", alignment=TA_CENTER)
    sml_left = ParagraphStyle("SmL", parent=styles["Normal"], fontSize=8, leading=9, alignment=TA_LEFT)

    elements = []

    # ── MÉTTRICAS GENERALES ────────────────────────────────
    if not df_sgmt.empty:
        imp_cols = [c for c in df_sgmt.columns if c.endswith("_Imp") and c != "Total_Imp"]
        df_sgmt["Total_Imp"] = df_sgmt[imp_cols].sum(axis=1)

    total_previsto = df_ud["horas_ud"].sum() if not df_ud.empty and "horas_ud" in df_ud.columns else 0
    total_impartido = df_sgmt["Total_Imp"].sum() if not df_sgmt.empty and "Total_Imp" in df_sgmt.columns else 0
    porcentaje = (total_impartido / total_previsto * 100) if total_previsto > 0 else 0

    h_real_total = 0
    h_sin_docencia = 0
    dias_semana_list = ["Lun", "Mar", "Mié", "Jue", "Vie"]
    for tri in ["1t", "2t", "3t"]:
        if f"ini_{tri}" in info_fechas and f"fin_{tri}" in info_fechas:
            ini = info_fechas[f"ini_{tri}"]
            fin = info_fechas[f"fin_{tri}"]
            curr = ini
            while curr <= fin:
                if curr.weekday() < 5:
                    f_str = curr.strftime("%d/%m/%Y")
                    if not calendar_notes.get(f"f_{f_str}"):
                        h_dia = horario.get(dias_semana_list[curr.weekday()], 0)
                        h_real_total += h_dia
                        if daily_ledger.get(f_str, {}).get("sin_docencia", False):
                            h_sin_docencia += h_dia
                curr += timedelta(days=1)
    perc_sin_docencia = (h_sin_docencia / h_real_total * 100) if h_real_total > 0 else 0

    metrics_text = f"<b>Horas Previstas:</b> {total_previsto} h   |   <b>Horas Impartidas:</b> {int(total_impartido)} h   |   <b>Progreso:</b> {porcentaje:.1f}%   |   <b>Sin docencia:</b> {perc_sin_docencia:.1f}%"
    elements.append(Paragraph(metrics_text, ParagraphStyle("Metrics", parent=norm, alignment=TA_LEFT, fontSize=11)))
    elements.append(Spacer(1, 15))

    # ── TABLA DE DATOS ──────────────────────────────────────
    meses_display = ["Sep", "Oct", "Nov", "Dic", "Ene", "Feb", "Mar", "Abr", "May", "Jun"]
    # Cod.UD, Ev., Clase Prv., Clase Imp. + 2 columnas (Prv/Imp) por mes
    col_widths = [1.5*cm, 0.9*cm, 1.1*cm, 1.1*cm] + [1.15*cm] * 20
    N_FIXED = 4

    table_data = []

    # Cabecera - fila 1: nombre de mes en una celda fusionada sobre sus dos
    # columnas Prv/Imp, igual que en la vista de la app (Planificación
    # mensual). Fila 2: etiquetas de columna.
    header_row1 = [""] * N_FIXED
    header_row2 = [
        Paragraph("<b>Cod.<br/>UD</b>", smlB),
        Paragraph("<b>Ev.</b>", smlB),
        Paragraph("<b>Clase<br/>Prv.</b>", sml),
        Paragraph("<b>Clase<br/>Imp.</b>", smlB),
    ]
    for m in meses_display:
        header_row1.append(Paragraph(f"<b>{m}</b>", smlB))
        header_row1.append("")
        header_row2.append(Paragraph("Prv.", sml))
        header_row2.append(Paragraph("<b>Imp.</b>", smlB))
    table_data.append(header_row1)
    table_data.append(header_row2)

    # Fila Sin docencia
    hoy_dt = datetime.now()
    total_incidencias = 0
    for d_str, entry in daily_ledger.items():
        try:
            d_obj = datetime.strptime(d_str, "%d/%m/%Y")
            if entry.get("sin_docencia") and d_obj.date() <= hoy_dt.date():
                total_incidencias += 1
        except: continue

    counts_meses = {}
    for m_short in meses_display:
        m_num = {"Sep":9, "Oct":10, "Nov":11, "Dic":12, "Ene":1, "Feb":2, "Mar":3, "Abr":4, "May":5, "Jun":6}[m_short]
        count_sin = 0
        for d_str, entry in daily_ledger.items():
            try:
                date_obj = datetime.strptime(d_str, "%d/%m/%Y")
                if date_obj.month == m_num and entry.get("sin_docencia") and date_obj.date() <= hoy_dt.date():
                    count_sin += 1
            except: continue
        counts_meses[m_short] = count_sin

    row_sin = [
        Paragraph("<b>Sin docencia</b>", smlB), "", "",
        Paragraph(f"<b><font color='#a36b00'>{total_incidencias}</font></b>", smlB) if total_incidencias > 0 else ""
    ]
    for m in meses_display:
        val = counts_meses[m]
        row_sin.append("") # Columna Prv vacía
        row_sin.append(Paragraph(f"<b><font color='#a36b00'>{val}</font></b>", smlB) if val > 0 else "")
    table_data.append(row_sin)

    # Filas de la tabla (por UD)
    ud_desc_map = {}
    if not df_ud.empty and "id_ud" in df_ud.columns:
        for _, ud_row in df_ud.iterrows():
            if ud_row.get("id_ud"):
                ud_desc_map[str(ud_row["id_ud"])] = str(ud_row.get("desc_ud", "") or "")

    if not df_sgmt.empty:
        for _, row in df_sgmt.iterrows():
            tot_i = row.get("Total_Imp", 0)
            u_p = row.get("horas_ud", 0)
            ev = row.get("ev", "")

            row_data = [
                Paragraph(f"<b>{row.get('id_ud', '')}</b>", smlB),
                Paragraph(f"{ev}ª" if ev else "", norm),
                Paragraph(f"{u_p}", norm),
                Paragraph(f"<b>{int(tot_i) if tot_i != 0 else ''}</b>", normB),
            ]
            for m in meses_display:
                val_p = row.get(f"{m}_Prv", 0)
                val_i = row.get(f"{m}_Imp", 0)

                str_p = str(int(val_p)) if val_p != 0 else ""
                str_i = str(int(val_i)) if val_i != 0 else ""

                row_data.append(Paragraph(str_p, norm))
                row_data.append(Paragraph(f"<b>{str_i}</b>" if str_i else "", normB))

            table_data.append(row_data)

    # ── ESTILOS DE TABLA ─────────────────────────────────────
    ts_cmds = [
        ("BACKGROUND",    (0, 0), (-1, 1), colors.HexColor("#f0f0f0")),
        ("TEXTCOLOR",     (0, 0), (-1, 1), colors.black),
        ("FONTNAME",      (0, 0), (-1, 1), "Helvetica-Bold"),
        ("LINEBELOW",     (0, 1), (-1, 1), 1.5, colors.HexColor("#222222")),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("BOX",           (0, 0), (-1, -1), 1.5, colors.HexColor("#222222")),
        ("GRID",          (0, 1), (-1, -1), 0.5, colors.HexColor("#bbbbbb")),
        ("TOPPADDING",    (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING",   (0, 0), (-1, -1), 2),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 2),
        ("SPAN",          (0, 2), (1, 2)), # Span "Sin docencia" row header
        ("BACKGROUND",    (0, 2), (-1, 2), colors.HexColor("#fffdf2")), # Fila sin docencia destacada
    ]

    # Fusionar el nombre de cada mes sobre sus 2 columnas (Prv/Imp) en la fila 1
    idx = N_FIXED
    for _ in meses_display:
        ts_cmds.append(("SPAN", (idx, 0), (idx + 1, 0)))
        ts_cmds.append(("ALIGN", (idx, 0), (idx + 1, 0), "CENTER"))
        idx += 2

    # Destacar columnas "Imp" (Clase Imp. y luego pares mensuales) y dibujar línea separadora gruesa
    ts_cmds.append(("BACKGROUND", (3, 1), (3, -1), colors.HexColor("#f8faff")))
    ts_cmds.append(("LINEAFTER", (3, 1), (3, -1), 2.0, colors.HexColor("#111111")))
    idx = N_FIXED + 1
    for _ in range(10):
        ts_cmds.append(("BACKGROUND", (idx, 1), (idx, -1), colors.HexColor("#f8fae6")))
        ts_cmds.append(("LINEAFTER", (idx, 1), (idx, -1), 2.0, colors.HexColor("#111111")))
        idx += 2

    ts = TableStyle(ts_cmds)

    tabla = Table(table_data, colWidths=col_widths, repeatRows=2)
    tabla.setStyle(ts)
    elements.append(tabla)

    # ── LEYENDA DE UD ─────────────────────────────────────
    if ud_desc_map:
        elements.append(Spacer(1, 14))
        for id_ud in sorted(ud_desc_map.keys()):
            desc = ud_desc_map[id_ud]
            texto = f"<b>{id_ud}</b> - {desc}" if desc else f"<b>{id_ud}</b>"
            elements.append(Paragraph(texto, sml_left))

    doc.build(elements)
    buffer.seek(0)
    return buffer


def generar_docx_planificacion(info_modulo, df_ud, df_sgmt, daily_ledger, horario, info_fechas, calendar_notes):
    """Versión .docx editable: una fila por UD, con Ev., horas previstas/
    impartidas totales, y Prv./Imp. por mes bajo una cabecera de 2 niveles
    (nombre de mes fusionado sobre sus columnas Prv/Imp), como en la app."""
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx_helpers import new_document, add_title, add_meta_line, doc_to_bytes

    doc = new_document(landscape=True)
    add_title(doc, "Planificación: horas previstas frente a impartidas", info_modulo.get("modulo", "Módulo"))
    add_meta_line(doc, f"{info_modulo.get('centro', '')} ({info_modulo.get('profesorado', '')})")

    if not df_sgmt.empty:
        imp_cols = [c for c in df_sgmt.columns if c.endswith("_Imp") and c != "Total_Imp"]
        df_sgmt = df_sgmt.copy()
        df_sgmt["Total_Imp"] = df_sgmt[imp_cols].sum(axis=1)

    total_previsto = df_ud["horas_ud"].sum() if not df_ud.empty and "horas_ud" in df_ud.columns else 0
    total_impartido = df_sgmt["Total_Imp"].sum() if not df_sgmt.empty and "Total_Imp" in df_sgmt.columns else 0
    porcentaje = (total_impartido / total_previsto * 100) if total_previsto > 0 else 0
    doc.add_paragraph(
        f"Horas previstas: {total_previsto} h    |    Horas impartidas: {int(total_impartido)} h    |    Progreso: {porcentaje:.1f}%"
    )

    def shade(cell, hex_color):
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shd)

    def set_cell(cell, text, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.CENTER):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(size)

    meses_display = ["Sep", "Oct", "Nov", "Dic", "Ene", "Feb", "Mar", "Abr", "May", "Jun"]
    N_FIXED = 4  # Cod. UD, Ev., H. Prv., H. Imp.
    n_cols = N_FIXED + 2 * len(meses_display)

    n_data_rows = len(df_sgmt) if not df_sgmt.empty else 0
    table = doc.add_table(rows=2 + n_data_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    HEADER_BG = "F0F0F0"
    fixed_headers = ["Cod. UD", "Ev.", "H. Prv.", "H. Imp."]
    for i, h in enumerate(fixed_headers):
        cell = table.rows[0].cells[i].merge(table.rows[1].cells[i])
        set_cell(cell, h, bold=True, size=8)
        shade(cell, HEADER_BG)

    for mi, m in enumerate(meses_display):
        c0 = N_FIXED + mi * 2
        mes_cell = table.rows[0].cells[c0].merge(table.rows[0].cells[c0 + 1])
        set_cell(mes_cell, m, bold=True, size=8)
        shade(mes_cell, HEADER_BG)
        for j, label in enumerate(("Prv.", "Imp.")):
            cell = table.rows[1].cells[c0 + j]
            set_cell(cell, label, bold=True, size=7)
            shade(cell, HEADER_BG)

    if not df_sgmt.empty:
        for r_idx, (_, row) in enumerate(df_sgmt.iterrows()):
            r = 2 + r_idx
            ev = row.get("ev", "")
            set_cell(table.rows[r].cells[0], row.get("id_ud", ""), bold=True, size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell(table.rows[r].cells[1], f"{ev}ª" if ev else "", size=8)
            set_cell(table.rows[r].cells[2], row.get("horas_ud", 0) or "", size=8)
            tot_i = row.get("Total_Imp", 0)
            set_cell(table.rows[r].cells[3], int(tot_i) if tot_i else "", bold=True, size=8)
            for mi, m in enumerate(meses_display):
                c0 = N_FIXED + mi * 2
                p = int(row.get(f"{m}_Prv", 0) or 0)
                i_ = int(row.get(f"{m}_Imp", 0) or 0)
                set_cell(table.rows[r].cells[c0], p if p else "", size=7)
                set_cell(table.rows[r].cells[c0 + 1], i_ if i_ else "", bold=True, size=7)

    col_widths = [Cm(2), Cm(1), Cm(1.5), Cm(1.5)] + [Cm(1.1)] * (2 * len(meses_display))
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

    doc.add_paragraph("")

    # ── Leyenda de UD ─────────────────────────────────────
    if not df_ud.empty and "id_ud" in df_ud.columns:
        for _, ud_row in df_ud.sort_values("id_ud").iterrows():
            id_ud = ud_row.get("id_ud", "")
            if not id_ud:
                continue
            desc = ud_row.get("desc_ud", "") or ""
            p = doc.add_paragraph()
            run_id = p.add_run(str(id_ud))
            run_id.bold = True
            if desc:
                p.add_run(f" - {desc}")

    return doc_to_bytes(doc)

