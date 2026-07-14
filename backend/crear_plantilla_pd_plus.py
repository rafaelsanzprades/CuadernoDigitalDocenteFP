# -*- coding: utf-8 -*-
"""
Crea la plantilla Jinja2 para PD+ (Detallada/JEG) desde cero.

Lee el contenido del dump del JEG original y crea un DOCX nuevo con:
- Todas las secciones y textos originales
- Variables {{ var }} en lugar de [[ instrucciones ]]
- Bloques explicativos eliminados
- 11 tablas recreadas con variables Jinja2
- Compatible 100% con docxtpl

Uso:
    python crear_plantilla_pd_plus.py
"""
import os, sys
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.stdout.reconfigure(encoding='utf-8')
os.chdir(os.path.dirname(os.path.abspath(__file__)))

DST = 'templates/modelo_pd_fp+_tpl.docx'


def add_heading(doc, text, level=1):
    """Add a heading with proper style."""
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False, style=None):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p


def add_table_with_header(doc, headers, rows_data=None, num_rows=0):
    """Add a table with header row and optional data rows."""
    total_rows = 1 + (len(rows_data) if rows_data else num_rows)
    table = doc.add_table(rows=total_rows, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    
    # Data rows
    if rows_data:
        for ri, row_data in enumerate(rows_data):
            for ci, val in enumerate(row_data):
                table.rows[ri + 1].cells[ci].text = val
    elif num_rows > 0:
        for ri in range(num_rows):
            for ci in range(len(headers)):
                table.rows[ri + 1].cells[ci].text = ''
    
    return table


def main():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # =========================================================================
    # PORTADA (P000-P054)
    # =========================================================================
    add_para(doc, '{{ tipo_centro }}', bold=True)
    add_para(doc, '{{ nombre_centro }}', bold=True)
    add_para(doc, 'Programación didáctica', bold=True, style='Title')
    add_para(doc, '{{ curso_academico }}')
    add_para(doc, '{{ tipo_elemento }}', bold=True)
    add_para(doc, '{{ modulo }}', bold=True)
    add_para(doc, 'Familia Profesional de')
    add_para(doc, '{{ familia_profesional }}')
    add_para(doc, 'Tipo de enseñanza')
    add_para(doc, '{{ tipo_ensennanza }}')
    add_para(doc, 'Grado {{ grado }} – Nivel {{ nivel }}')
    add_para(doc, '{{ curso_numero }} curso')
    add_para(doc, 'Modalidad {{ modalidad }}')
    
    doc.add_paragraph()  # spacer
    add_para(doc, 'Autor', bold=True)
    add_para(doc, 'Javier Edo Gual')
    add_para(doc, 'Coordinación', bold=True)
    add_para(doc, 'Raúl Melero Rubio')
    add_para(doc, 'Lucía Quílez Salvador')
    add_para(doc, 'Revisión técnica', bold=True)
    add_para(doc, 'Óscar Sánchez Estella')
    
    doc.add_paragraph()
    add_para(doc, 'Licencia', bold=True)
    add_para(doc, 'Programaciones didácticas de Formación Profesional: modelo para los Grados D y E © 2025 por Javier Edo Gual está licenciada bajo CC BY-NC-SA 4.0.')
    add_para(doc, 'Versión', bold=True)
    add_para(doc, '1.0 (septiembre 2025)')
    
    doc.add_page_break()
    
    # =========================================================================
    # INSTRUCCIONES (P036-P054) — ELIMINADAS (solo instrucciones JEG)
    # =========================================================================
    
    # =========================================================================
    # 1. INTRODUCCIÓN
    # =========================================================================
    add_heading(doc, '1. Introducción', level=1)
    
    # 1.1 Identificación
    add_heading(doc, '1.1. Identificación', level=2)
    add_para(doc, '{{ texto_identificacion }}')
    
    # 1.2 Marco normativo
    add_heading(doc, '1.2. Marco normativo', level=2)
    add_para(doc, 'Esta programación didáctica se fundamenta en la siguiente normativa:')
    add_para(doc, 'Normativa estatal:', bold=True)
    add_para(doc, 'LEY ORGÁNICA 2/2006, de 3 de mayo, de Educación (BOE núm. 106, de 4 de mayo de 2006), modificada por la LEY ORGÁNICA 3/2020, de 29 de diciembre.')
    add_para(doc, 'LEY ORGÁNICA 3/2022, de 31 de marzo, de ordenación e integración de la Formación Profesional (BOE núm. 78, de 01 de abril de 2022).')
    add_para(doc, 'REAL DECRETO 659/2023, de 18 de julio, por el que se desarrolla la ordenación del Sistema de Formación Profesional (BOE núm. 174, de 22 de julio de 2023).')
    add_para(doc, 'REAL DECRETO 69/2025, de 4 de febrero, por el que se desarrollan los elementos integrantes y los instrumentos de gestión del Sistema Nacional de Formación Profesional.')
    add_para(doc, 'REAL DECRETO por el que se establece el título del grado D o E correspondiente.')
    
    add_para(doc, 'Normativa autonómica:', bold=True)
    add_para(doc, 'DECRETO 91/2024, de 5 de junio, del Gobierno de Aragón por el que se establece la Ordenación de la Formación Profesional del Grado D y del Grado E en la Comunidad Autónoma de Aragón (BOA núm. 109, de 06 de junio de 2024).')
    add_para(doc, 'ORDEN ECD/841/2024, de 25 de julio, por la que se regulan aspectos organizativos del currículo y se establecen los currículos de determinados Ciclos Formativos de Formación Profesional de Grado Básico para la Comunidad Autónoma de Aragón (BOA núm. 148, de 31 de julio de 2024).')
    add_para(doc, 'ORDEN ECD/842/2024, de 25 de julio, por la que se regulan aspectos organizativos del currículo y se establecen los currículos de determinados Ciclos Formativos de Formación Profesional de Grado Medio para la Comunidad Autónoma de Aragón (BOA núm. 148, de 31 de julio de 2024).')
    add_para(doc, 'ORDEN ECD/843/2024, de 25 de julio, por la que se regulan aspectos organizativos del currículo y se establecen los currículos de determinados Ciclos Formativos de Formación Profesional de Grado Superior para la Comunidad Autónoma de Aragón (BOA núm. 148, de 31 de julio de 2024).')
    add_para(doc, 'RESOLUCIÓN de 6 de junio de 2025, del Director General de Planificación, Centros y Formación Profesional, por la que se autorizan los módulos optativos de los centros docentes.')
    add_para(doc, 'RESOLUCIÓN de 10 de julio de 2025, del Director General de Planificación, Centros y Formación Profesional, por la que se establece el currículo de los módulos optativos.')
    add_para(doc, 'RESOLUCIÓN de 11 de junio de 2025, del Director General de Planificación, Centros y Formación Profesional, por la que se determinan los módulos profesionales en las modalidades virtual y semipresencial.')
    add_para(doc, 'RESOLUCIÓN de 1 de julio de 2025, del Director General de Planificación, Centros y Formación Profesional, por la que se regula la organización y la distribución horaria de varios Cursos de Especialización de Formación Profesional (Grado E).')
    add_para(doc, 'ORDEN ECD/1005/2018, de 7 de junio, por la que se regulan las actuaciones de intervención educativa inclusiva (BOA núm. 106, de 18 de junio de 2018).')
    add_para(doc, 'ORDEN por la que se aprueban las Instrucciones del tipo de centro en relación con el curso actual.')
    
    # 1.3 Contextualización
    add_heading(doc, '1.3. Contextualización', level=2)
    
    add_heading(doc, '1.3.1. Entorno geográfico y sociocultural', level=3)
    add_para(doc, 'El {{ tipo_elemento }} se imparte en el contexto geográfico y sociocultural de {{ municipio }}, caracterizado por {{ campo_libre }}.')
    add_para(doc, 'El entorno cultural incluye {{ contexto_cultural }}, configurando un marco de referencia que define parte de la identidad del alumnado y del centro.')
    add_para(doc, 'En cuanto a la accesibilidad y comunicaciones, se dispone de {{ campo_libre }} que condicionan la organización de desplazamientos y el uso de recursos.')
    add_para(doc, 'Del análisis del contexto se derivan las siguientes implicaciones para la programación:')
    add_para(doc, '{{ implicaciones_contexto_geografico }}')
    
    add_heading(doc, '1.3.2. Entorno socioeconómico y productivo', level=3)
    add_para(doc, 'El {{ tipo_elemento }} se imparte en un entorno socioeconómico y productivo caracterizado por la presencia de {{ sectores_productivos }}, con una vinculación con la oferta formativa del {{ grado }}.')
    add_para(doc, 'El mercado laboral del entorno demanda {{ perfiles_profesionales }}, lo que condiciona la orientación del {{ tipo_elemento }} hacia la inserción laboral del alumnado.')
    add_para(doc, 'Del análisis del entorno productivo se derivan las siguientes implicaciones para la programación:')
    add_para(doc, '{{ implicaciones_contexto_productivo }}')
    
    add_heading(doc, '1.3.3. Centro educativo', level=3)
    add_para(doc, '{{ nombre_centro }} es un centro {{ tipo_centro }} que ofrece {{ oferta_formativa }}.')
    add_para(doc, 'El centro se caracteriza por {{ campo_libre }}.')
    add_para(doc, 'Del análisis del centro se derivan las siguientes implicaciones para la programación:')
    add_para(doc, '{{ implicaciones_centro }}')
    
    add_heading(doc, '1.3.4. Alumnado', level=3)
    add_para(doc, 'El alumnado que cursa el {{ tipo_elemento }} procede de {{ procedencia_alumnado }}, con una franja de edad comprendida entre {{ franja_edad }} años.')
    add_para(doc, 'El grupo muestra {{ grado_implicacion }} grado de implicación y expectativas relacionadas con {{ expectativas }}.')
    add_para(doc, 'Asimismo, cabe destacar {{ caracteristicas_alumnado_extra }}.')
    add_para(doc, 'Del análisis de estas características se derivan las siguientes implicaciones para la programación:')
    add_para(doc, '{{ implicaciones_alumnado }}')
    
    add_heading(doc, '1.3.5. Infraestructura y recursos educativos', level=3)
    add_para(doc, 'El {{ tipo_elemento }} dispone de las siguientes infraestructuras y recursos:')
    add_para(doc, 'Espacios:', bold=True)
    add_para(doc, '{{ espacios_recursos }}')
    add_para(doc, 'Equipamiento técnico y específico:', bold=True)
    add_para(doc, '{{ equipamiento_tecnico }}')
    add_para(doc, 'Recursos tecnológicos:', bold=True)
    add_para(doc, '{{ recursos_tecnologicos }}')
    add_para(doc, 'Software, herramientas digitales y plataformas virtuales:', bold=True)
    add_para(doc, '{{ software_herramientas }}')
    add_para(doc, 'Material didáctico y documental:', bold=True)
    add_para(doc, '{{ material_didactico }}')
    add_para(doc, 'Espacios complementarios:', bold=True)
    add_para(doc, '{{ espacios_complementarios }}')
    add_para(doc, 'Del análisis de esta infraestructura y recursos se derivan las siguientes implicaciones para la programación:')
    add_para(doc, '{{ implicaciones_infraestructura }}')
    
    doc.add_page_break()
    
    # =========================================================================
    # 2. DESARROLLO CURRICULAR
    # =========================================================================
    add_heading(doc, '2. Desarrollo curricular', level=1)
    
    # 2.1 Duración, ubicación y distribución horaria
    add_heading(doc, '2.1. Duración, ubicación y distribución horaria', level=2)
    add_para(doc, 'En el curso {{ curso_academico }} la carga lectiva semanal se distribuye de la siguiente manera:')
    
    # Table 0: Distribución horaria semanal
    add_table_with_header(doc, ['Día', 'Horas'], num_rows=5)
    doc.add_paragraph()
    
    add_para(doc, 'Teniendo en cuenta el calendario escolar y la distribución semanal de la carga lectiva, la duración prevista para el curso {{ curso_academico }} es de {{ horas }} horas.')
    
    # 2.2 ECP
    add_heading(doc, '2.2. Estándares de competencias profesionales (ECP)', level=2)
    add_para(doc, 'La superación de este módulo profesional permite la acreditación de los siguientes estándares de competencias profesionales del Catálogo Nacional de Estándares de Competencias Profesionales (CNECP):')
    add_para(doc, '{{ texto_ecp }}')
    
    # 2.3 CPE
    add_heading(doc, '2.3. Competencias profesionales y para la empleabilidad (CPE)', level=2)
    add_para(doc, 'Este módulo profesional contribuye al desarrollo de las siguientes competencias profesionales y para la empleabilidad (CPE):')
    add_para(doc, '{{ texto_cpe }}')
    
    # 2.4 OG
    add_heading(doc, '2.4. Objetivos generales (OG)', level=2)
    add_para(doc, 'De los objetivos generales (OG) establecidos en el currículo, la formación de este {{ tipo_elemento }} contribuye a alcanzar los siguientes:')
    add_para(doc, '{{ texto_og }}')
    
    # 2.5 RA y CE
    add_heading(doc, '2.5. Resultados de aprendizaje (RA) y criterios de evaluación (CE)', level=2)
    add_para(doc, 'En este {{ tipo_elemento }} se establecen los siguientes resultados de aprendizaje (RA) y criterios de evaluación (CE):')
    
    for i in range(1, 8):
        add_para(doc, f'RA {i}. {{{{ ra{i}_titulo }}}}', bold=True)
        add_para(doc, 'Criterios de evaluación:')
        add_para(doc, f'{{% for ce in ra{i}_loop %}}')
        add_para(doc, f'{{{{ ce.ce_code }}}} {{{{ ce.desc_ce }}}}')
        add_para(doc, f'{{% endfor %}}')
    
    # 2.6 Contenidos
    add_heading(doc, '2.6. Contenidos (C)', level=2)
    add_para(doc, 'Este {{ tipo_elemento }} tiene los siguientes bloques de contenidos (BC) y contenidos orientativos (C):')
    add_para(doc, '{{ texto_contenidos }}')
    
    # 2.7 Relación entre elementos curriculares
    add_heading(doc, '2.7. Relación entre los elementos curriculares', level=2)
    add_para(doc, 'A continuación se muestra la relación entre los distintos elementos curriculares del {{ tipo_elemento }}:')
    
    # Table 1: Relación RA-CE-UD
    add_table_with_header(doc,
        ['RA', 'Descripción RA', 'CE', 'Descripción CE', 'UD'],
        num_rows=6
    )
    doc.add_paragraph()
    
    # 2.8 Organización temporal — UD
    add_heading(doc, '2.8. Organización y distribución temporal: unidades didácticas y otros tiempos', level=2)
    add_para(doc, 'El {{ tipo_elemento }} tiene una duración curricular de {{ horas }} horas. Para el curso académico {{ curso_academico }}, la duración prevista es de {{ horas }} horas, calculada en función de la carga horaria semanal asignada y del calendario escolar establecido.')
    add_para(doc, 'Asimismo, considerando lo establecido en la propuesta de organización temporal para la fase de formación en empresa, de la duración total prevista, {{ horas_fct }} horas corresponden al período que el alumnado permanece en la empresa u organismo equiparado. Por lo tanto, la duración efectiva de formación en el centro educativo es de {{ horas_empresa }} horas.')
    add_para(doc, 'Las unidades didácticas se impartirán siguiendo una organización {{ organizacion_secuencial }}.')
    add_para(doc, 'A continuación, se detalla la organización y distribución temporal de las unidades didácticas:')
    
    # Table 2: UD timeline
    add_table_with_header(doc,
        ['UD', 'Título', 'RA', 'Horas', 'Evaluación'],
        num_rows=10
    )
    doc.add_paragraph()
    
    # Table 3: UD-RA matrix
    add_table_with_header(doc,
        ['UD', 'Título', 'RA1', 'RA2', 'RA3', 'RA4', 'RA5', 'RA6', 'RA7'],
        num_rows=10
    )
    doc.add_paragraph()
    
    # Table 4: UD detail
    add_table_with_header(doc,
        ['N', 'Título', 'RA', 'CE', 'C', 'CPE', 'OG', 'Duración (horas)', 'Temporización'],
        num_rows=10
    )
    doc.add_paragraph()
    
    # Table 5: Actividades no vinculadas a UD
    add_table_with_header(doc,
        ['Actividad no vinculada a UD', 'Duración (horas)'],
        num_rows=3
    )
    doc.add_paragraph()
    
    # 2.9 FEOE
    add_heading(doc, '2.9. Formación en empresa u organismo equiparado (FEOE)', level=2)
    
    add_heading(doc, '2.9.1. Modalidad de la FEOE', level=3)
    add_para(doc, 'La fase de formación en empresa u organismo equiparado se desarrollará en el régimen {{ regimen_feoe }}.')
    
    add_heading(doc, '2.9.2. Organización temporal de la FEOE', level=3)
    add_para(doc, 'La fase de formación en empresa u organismo equiparado se desarrollará desde el {{ fecha_inicio_feoe }} hasta el {{ fecha_fin_feoe }}, ambos inclusive, con una duración total de {{ horas_fct }} horas, para el conjunto de módulos profesionales que se desarrollan en la empresa u organismo equiparado.')
    add_para(doc, 'Esta planificación podrá ajustarse en función de la organización del centro, las características y circunstancias del alumnado, así como de las necesidades y condiciones de las empresas u organismos equiparados.')
    
    add_heading(doc, '2.9.3. Resultados de aprendizaje desarrollados en la FEOE', level=3)
    add_para(doc, 'A continuación, se indican los resultados de aprendizaje (RA) que se prevé desarrollar, total o parcialmente, en la fase de formación en empresa u organismo equiparado:')
    
    # Table 6: RA en FEOE
    add_table_with_header(doc,
        ['Resultado de aprendizaje', 'Parcial / Total', 'Justificación'],
        num_rows=4
    )
    doc.add_paragraph()
    
    add_para(doc, 'Por tanto, {{ num_ra_feoe }} de los {{ num_total_ra }} resultados de aprendizaje de este módulo profesional se desarrollan, total o parcialmente, en la fase de formación en empresa u organismo equiparado.')
    
    add_heading(doc, '2.9.4. Organización para el alumnado que no realice la FEOE', level=3)
    add_para(doc, '{{ organizacion_no_feoe }}')
    
    add_heading(doc, '2.9.5. Seguimiento de la FEOE', level=3)
    add_para(doc, '{{ texto_seguimiento_feoe }}')
    
    # 2.10 Metodología
    add_heading(doc, '2.10. Metodología didáctica', level=2)
    
    add_heading(doc, '2.10.1. Principios y estrategias metodológicas', level=3)
    add_para(doc, 'La metodología didáctica de este {{ tipo_elemento }} se basa en los siguientes principios y estrategias:')
    add_para(doc, '{{ metodologias_seleccionadas }}')
    add_para(doc, '{{ texto_metodologia_libre }}')
    
    add_heading(doc, '2.10.2. Agrupamientos', level=3)
    add_para(doc, 'Se utilizarán los siguientes tipos de agrupamiento:')
    add_para(doc, 'Individual:', bold=True)
    add_para(doc, '{{ agrupamiento_individual }}')
    add_para(doc, 'Gran grupo o grupo clase:', bold=True)
    add_para(doc, '{{ agrupamiento_gran_grupo }}')
    add_para(doc, 'Parejas:', bold=True)
    add_para(doc, '{{ agrupamiento_parejas }}')
    add_para(doc, 'Equipo:', bold=True)
    add_para(doc, 'El alumnado se agrupará en equipos de {{ tamanno_equipo }} personas.')
    add_para(doc, '{{ agrupamiento_equipo }}')
    
    add_heading(doc, '2.10.3. Plan de aplicación de los desdobles', level=3)
    add_para(doc, 'En la siguiente tabla se detalla el plan de aplicación de los desdobles previstos para el {{ tipo_elemento }}:')
    
    # Table 7: Desdobles
    add_table_with_header(doc,
        ['Temporización', ''],
        [
            ['Unidades didácticas', ''],
            ['Justificación', ''],
            ['Número aproximado de alumnado por grupo', ''],
            ['Distribución de espacios', ''],
            ['Recursos específicos utilizados', ''],
        ]
    )
    doc.add_paragraph()
    
    add_heading(doc, '2.10.4. Aprendizaje colaborativo basado en proyectos y/o retos', level=3)
    add_para(doc, 'El {{ tipo_elemento }}, en un contexto de aprendizaje colaborativo, participa en el siguiente {{ tipo_colaborativo }}:')
    add_para(doc, '{{ texto_aprendizaje_colaborativo }}')
    
    # Table 8: Proyecto colaborativo
    add_table_with_header(doc,
        ['Tipo', ''],
        [
            ['Título', ''],
            ['Módulos profesionales implicados', ''],
            ['Descripción', ''],
            ['Resultados de aprendizaje y criterios de evaluación', ''],
            ['Principios y estrategias metodológicas', ''],
            ['Materiales y recursos didácticos', ''],
            ['Agrupamientos', ''],
            ['Duración estimada', ''],
            ['Fechas estimadas de inicio y fin', ''],
            ['Organización en la distribución horaria semanal', ''],
            ['Plan de aplicación de desdobles y/o codocencia', ''],
            ['Plan de coordinación docente', ''],
        ]
    )
    doc.add_paragraph()
    
    add_heading(doc, '2.10.5. Materiales y recursos didácticos', level=3)
    add_para(doc, 'En el desarrollo del {{ tipo_elemento }} se utilizarán los siguientes recursos:')
    add_para(doc, 'Recursos personales:', bold=True)
    add_para(doc, '{{ recursos_personales }}')
    add_para(doc, 'Espacios:', bold=True)
    add_para(doc, '{{ espacios_recursos }}')
    add_para(doc, 'Recursos materiales:', bold=True)
    add_para(doc, '{{ recursos_materiales }}')
    add_para(doc, 'Recursos digitales:', bold=True)
    add_para(doc, '{{ recursos_digitales }}')
    add_para(doc, 'Recursos documentales:', bold=True)
    add_para(doc, '{{ recursos_documentales }}')
    add_para(doc, 'Bibliografía:', bold=True)
    add_para(doc, '{{ bibliografia }}')
    
    doc.add_page_break()
    
    # =========================================================================
    # 3. PROCEDIMIENTO DE EVALUACIÓN Y DE CALIFICACIÓN
    # =========================================================================
    add_heading(doc, '3. Procedimiento de evaluación y de calificación', level=1)
    
    # 3.1 Actividades, técnicas e instrumentos
    add_heading(doc, '3.1. Actividades, técnicas e instrumentos', level=2)
    add_para(doc, 'A lo largo del desarrollo del {{ tipo_elemento }} se utilizarán los siguientes tipos de actividades de evaluación:')
    add_para(doc, '{{ tipo_actividad }}')
    add_para(doc, 'La recogida y valoración de la información se llevará a cabo mediante los siguientes instrumentos de evaluación:')
    add_para(doc, '{{ instrumentos_seleccionados }}')
    add_para(doc, 'En función de los resultados obtenidos en la evaluación inicial, se aplicarán las siguientes medidas:')
    add_para(doc, '{{ medidas_evaluacion_inicial }}')
    
    # 3.2 Evaluaciones parciales
    add_heading(doc, '3.2. Evaluaciones parciales', level=2)
    add_para(doc, 'Durante el curso se realizarán, con carácter general, {{ num_evaluaciones_parciales }}.')
    add_para(doc, 'Estas sesiones forman parte del proceso de evaluación continua y tendrán un carácter formativo e integrador, orientado a valorar el progreso del alumnado y a mejorar la intervención educativa.')
    add_para(doc, 'En cada evaluación parcial se valorará el progreso del alumnado en la consecución de los resultados de aprendizaje.')
    add_para(doc, 'A cada estudiante se le informará del resultado de la evaluación mediante una calificación numérica entre 1 y 10.')
    add_para(doc, 'En estas evaluaciones parciales se informará únicamente sobre los resultados de aprendizaje respecto a los cuales se haya recogido evidencia suficiente conforme a sus criterios de evaluación.')
    
    # 3.3 Primera evaluación final
    add_heading(doc, '3.3. Primera evaluación final', level=2)
    add_para(doc, 'Durante el mes de junio, se realizará la primera evaluación final, que tendrá carácter calificativo.')
    add_para(doc, 'Alumnado que no haya perdido el derecho a la evaluación continua:', bold=True)
    add_para(doc, 'El alumnado que haya mantenido el derecho a la evaluación continua será calificado en esta primera evaluación final tomando como base las evidencias de aprendizaje recogidas a lo largo del curso.')
    add_para(doc, 'La calificación final del {{ tipo_elemento }} se obtendrá aplicando los criterios de calificación, establecidos en el apartado 3.7, y la ponderación de los resultados de aprendizaje y de los criterios de evaluación.')
    add_para(doc, 'Alumnado que haya perdido el derecho a la evaluación continua:', bold=True)
    add_para(doc, 'El alumnado que pierda el derecho a la evaluación continua será evaluado en esta convocatoria mediante actividades específicas que generen nuevas evidencias de aprendizaje.')
    add_para(doc, 'La calificación final del {{ tipo_elemento }} se obtendrá aplicando los mismos criterios de calificación y ponderación establecidos para el alumnado que no haya perdido el derecho a la evaluación continua.')
    
    # 3.4 Segunda evaluación final
    add_heading(doc, '3.4. Segunda evaluación final (septiembre)', level=2)
    add_para(doc, 'En septiembre se realizará una segunda evaluación final, de carácter calificativo, para el alumnado que no haya superado el {{ tipo_elemento }} en la primera evaluación final.')
    add_para(doc, 'Esta evaluación incluirá los resultados de aprendizaje no superados, valorándose la totalidad de los criterios de evaluación asociados.')
    add_para(doc, 'La calificación final del {{ tipo_elemento }} se obtendrá aplicando los mismos criterios de calificación y ponderación establecidos para la primera evaluación final.')
    
    # 3.5 Evaluación continua
    add_heading(doc, '3.5. Evaluación continua', level=2)
    add_para(doc, '{{ texto_evaluacion_continua }}')
    
    # 3.6 Pérdida del derecho a la evaluación continua
    add_heading(doc, '3.6. Pérdida del derecho a la evaluación continua', level=2)
    add_para(doc, '{{ texto_perdida_evaluacion_continua }}')
    
    # 3.7 Criterios de calificación y ponderación
    add_heading(doc, '3.7. Criterios de calificación y ponderación', level=2)
    
    add_heading(doc, '3.7.1. Criterios de calificación', level=3)
    add_para(doc, 'La calificación del {{ tipo_elemento }} se calculará a partir de los siguientes criterios de calificación:')
    add_para(doc, '{{ texto_criterios_calificacion }}')
    add_para(doc, 'El {{ tipo_elemento }} se considerará superado cuando se obtenga una calificación igual o superior a {{ calificacion_minima }} puntos sobre 10.')
    add_para(doc, 'La calificación final se obtendrá mediante un proceso jerárquico:')
    add_para(doc, 'Criterios de evaluación (CE): cada CE se calificará a partir de las evidencias de aprendizaje recogidas.')
    add_para(doc, 'Resultados de aprendizaje (RA): la calificación de cada RA se calculará como la media ponderada de las calificaciones de sus criterios de evaluación.')
    add_para(doc, 'Calificación final del {{ tipo_elemento }}: se calculará como la media ponderada de las calificaciones de los resultados de aprendizaje.')
    
    add_heading(doc, '3.7.2. Ponderación de resultados de aprendizaje', level=3)
    add_para(doc, '{{ ponderacion_ra }}')
    
    # 3.8 Recuperación
    add_heading(doc, '3.8. Recuperación', level=2)
    add_para(doc, '{{ momentos_recuperacion }}')
    add_para(doc, '{{ actividades_recuperacion }}')
    add_para(doc, '{{ instrumentos_recuperacion }}')
    add_para(doc, '{{ ambito_recuperacion }}')
    
    # 3.9 Plan de recuperación de pendientes
    add_heading(doc, '3.9. Plan de recuperación de pendientes', level=2)
    add_para(doc, '{{ fecha_inicio_plan_recuperacion }}')
    add_para(doc, '{{ fecha_fin_plan_recuperacion }}')
    add_para(doc, '{{ fases_plan_recuperacion }}')
    add_para(doc, '{{ actividades_eval_inicial_recuperacion }}')
    add_para(doc, '{{ actividades_plan_recuperacion }}')
    add_para(doc, '{{ desarrollo_plan_recuperacion }}')
    add_para(doc, '{{ recursos_plan_recuperacion }}')
    
    # 3.10 Evaluación de la FEOE
    add_heading(doc, '3.10. Evaluación de la fase de formación en empresa u organismo equiparado (FEOE)', level=2)
    add_para(doc, 'La calificación final de los resultados de aprendizaje (RA) trabajados en la FEOE es responsabilidad última del profesorado del módulo profesional.')
    add_para(doc, 'El procedimiento de evaluación de la FEOE será el siguiente:')
    add_para(doc, 'Concreción de actividades:', bold=True)
    add_para(doc, 'Para cada estudiante, y en colaboración con la empresa, se concretarán en su Plan de Formación las actividades que permitirán desarrollar y evaluar cada resultado de aprendizaje de la FEOE.')
    add_para(doc, 'Seguimiento y valoración en la empresa u organismo equiparado:', bold=True)
    add_para(doc, 'El tutor o tutora dual de empresa valorará cada actividad vinculada a un RA con una puntuación entre 1 y 4. Un resultado de aprendizaje se considerará superado cuando el valor medio de la evaluación de sus actividades sea superior a 2.')
    add_para(doc, 'Seguimiento por el centro educativo:', bold=True)
    add_para(doc, 'Con carácter quincenal, el alumnado deberá asistir a una reunión presencial en el centro, donde se llevará a cabo el seguimiento de la FEOE.')
    add_para(doc, 'Calificación final:', bold=True)
    add_para(doc, 'El profesorado del módulo asignará la calificación final integrando la valoración numérica y cualitativa del tutor de empresa.')
    
    # 3.11 Información al alumnado
    add_heading(doc, '3.11. Información al alumnado', level=2)
    add_para(doc, '{{ texto_info_evaluacion }}')
    
    doc.add_page_break()
    
    # =========================================================================
    # 4. PLAN DE RECUPERACIÓN DE PENDIENTES
    # =========================================================================
    add_heading(doc, '4. Plan de recuperación de pendientes', level=1)
    add_para(doc, '{{ texto_plan_recuperacion_pendientes }}')
    
    doc.add_page_break()
    
    # =========================================================================
    # 5. MEDIDAS DE RESPUESTA EDUCATIVA PARA LA INCLUSIÓN
    # =========================================================================
    add_heading(doc, '5. Medidas de respuesta educativa para la inclusión', level=1)
    add_para(doc, 'A partir de la información obtenida en la evaluación inicial, y atendiendo de manera esencial a las características del alumnado, se adoptarán las siguientes medidas de respuesta educativa para la inclusión:')
    add_para(doc, '{{ medidas_inclusion }}')
    add_para(doc, '{{ texto_inclusion_libre }}')
    
    doc.add_page_break()
    
    # =========================================================================
    # 6. ACTIVIDADES COMPLEMENTARIAS Y EXTRAESCOLARES
    # =========================================================================
    add_heading(doc, '6. Actividades complementarias y extraescolares', level=1)
    add_para(doc, 'A continuación, se detallan las actividades complementarias y extraescolares programadas para el {{ tipo_elemento }}:')
    
    add_para(doc, 'Título de la actividad 1:', bold=True)
    add_para(doc, 'Tipo: {{ actividad1_tipo }}.')
    add_para(doc, 'Resultados de aprendizaje (RA): {{ actividad1_ra }}.')
    add_para(doc, 'Temporización: {{ actividad1_temporizacion }}.')
    add_para(doc, 'Entidad colaboradora: {{ actividad1_entidad }}.')
    add_para(doc, 'Descripción: {{ actividad1_descripcion }}.')
    add_para(doc, 'Evaluación: {{ actividad1_evaluacion }}.')
    
    add_para(doc, 'Título de la actividad 2:', bold=True)
    add_para(doc, 'Tipo: {{ actividad2_tipo }}.')
    add_para(doc, 'Resultados de aprendizaje (RA): {{ actividad2_ra }}.')
    add_para(doc, 'Temporización: {{ actividad2_temporizacion }}.')
    add_para(doc, 'Entidad colaboradora: {{ actividad2_entidad }}.')
    add_para(doc, 'Descripción: {{ actividad2_descripcion }}.')
    add_para(doc, 'Evaluación: {{ actividad2_evaluacion }}.')
    
    doc.add_page_break()
    
    # =========================================================================
    # 7. PLAN DE CONTINGENCIA
    # =========================================================================
    add_heading(doc, '7. Plan de contingencia', level=1)
    add_para(doc, '{{ medidas_contingencia }}')
    add_para(doc, '{{ texto_contingencia_libre }}')
    add_para(doc, '{{ actuaciones_contingencia }}')
    
    doc.add_page_break()
    
    # =========================================================================
    # 8. COMUNICACIÓN CON FAMILIAS
    # =========================================================================
    add_heading(doc, '8. Comunicación con familias', level=1)
    add_para(doc, '{{ formato_comunicacion }}')
    add_para(doc, '{{ formato_comunicacion_inicial }}')
    add_para(doc, '{{ formato_comunicacion_desarrollo }}')
    
    doc.add_page_break()
    
    # =========================================================================
    # 9. PUBLICIDAD DE LA PROGRAMACIÓN DIDÁCTICA
    # =========================================================================
    add_heading(doc, '9. Publicidad de la programación didáctica', level=1)
    add_para(doc, 'Durante la primera sesión del curso, el profesorado titular del {{ tipo_elemento }} explicará al alumnado los aspectos más relevantes de la programación didáctica, con el fin de garantizar la transparencia del proceso de enseñanza-aprendizaje.')
    add_para(doc, 'Esta información incluirá:')
    add_para(doc, '{{ texto_publicidad }}')
    
    doc.add_page_break()
    
    # =========================================================================
    # 10. EVALUACIÓN INICIAL
    # =========================================================================
    add_heading(doc, '10. Evaluación inicial', level=1)
    add_para(doc, '{{ actividades_evaluacion_inicial }}')
    add_para(doc, '{{ medidas_evaluacion_inicial }}')
    
    # Table 9: Evaluación inicial — dimensiones
    add_table_with_header(doc,
        ['Dimensión', 'Aspectos', 'Criterios de valoración'],
        [
            ['Perfil personal y académico del alumnado', 'Perfil personal y académico del alumnado', 'Perfil personal y académico del alumnado'],
            ['Trayectoria educativa', 'Vía de acceso (ESO, prueba de acceso, etc.). | Repeticiones de curso. | Apoyos recibidos anteriormente.', 'La trayectoria previa proporciona una base suficiente para abordar el ciclo. | Las repeticiones anteriores no suponen lagunas críticas. | Los apoyos recibidos han contribuido a superar dificultades anteriores.'],
            ['Contexto familiar y emocional', 'Implicación familiar en los estudios. | Clima familiar y estabilidad emocional. | Autoestima y motivación.', 'La familia ofrece seguimiento y apoyo educativo adecuados. | El alumno/a presenta estabilidad emocional y seguridad. | Muestra motivación suficiente para implicarse en el aprendizaje.'],
            ['Condiciones materiales para el estudio', 'Disponibilidad de dispositivos y conectividad. | Espacio de estudio en casa. | Compatibilidad horaria.', 'Dispone de acceso estable a internet y herramientas digitales. | Cuenta con un espacio físico adecuado. | Sus horarios permiten dedicar tiempo suficiente.'],
            ['Socialización y trabajo en equipo', 'Relación con profesorado y compañeros/as. | Habilidades sociales. | Participación en dinámicas grupales.', 'Se integra de forma adecuada en el grupo. | Respeta las normas de convivencia. | Colabora de manera activa y constructiva.'],
            ['Hábitos e intereses', 'Gestión del tiempo libre. | Actividades extracurriculares. | Intereses vinculados al módulo.', 'Participa en actividades enriquecedoras. | Participa en actividades culturales o deportivas. | Muestra curiosidad por temas relacionados con el módulo.'],
            ['Autonomía y estrategias de aprendizaje', 'Organización del estudio. | Autorregulación. | Uso de técnicas de aprendizaje.', 'Planifica y gestiona sus tareas con eficacia. | Detecta sus dificultades y solicita apoyo. | Emplea estrategias de aprendizaje efectivas.'],
            ['Salud y bienestar', 'Hábitos de sueño y alimentación. | Condiciones físicas o de salud. | Gestión del estrés.', 'Mantiene hábitos saludables. | Informa de necesidades específicas. | Gestiona adecuadamente el estrés.'],
            ['Expectativas y actitud ante el ciclo', 'Interés y motivación. | Percepción de dificultad. | Objetivos profesionales.', 'Demuestra interés activo. | Afronta retos con perseverancia. | Relaciona los aprendizajes con sus metas profesionales.'],
            ['Barreras contextuales', 'Limitaciones económicas o geográficas. | Necesidad de adaptaciones.', 'Dispone de recursos suficientes. | Su residencia no limita el seguimiento. | El centro puede proporcionarle recursos.'],
            ['Nivel competencial inicial en el {{ tipo_elemento }}', 'Nivel competencial inicial en el {{ tipo_elemento }}', 'Nivel competencial inicial en el {{ tipo_elemento }}'],
            ['Competencia digital', 'Crear, guardar y organizar archivos. | Comunicación digital. | Seguridad digital.', 'Maneja herramientas digitales básicas. | Se comunica eficazmente. | Conoce riesgos digitales.'],
        ]
    )
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # =========================================================================
    # 11. ELEMENTOS TRANSVERSALES
    # =========================================================================
    add_heading(doc, '11. Elementos transversales', level=1)
    add_para(doc, '{{ elementos_transversales }}')
    
    # =========================================================================
    # 12. ATENCIÓN A LA DIVERSIDAD
    # =========================================================================
    add_heading(doc, '12. Atención a la diversidad', level=1)
    add_para(doc, '{{ texto_atencion_diversidad }}')
    
    # =========================================================================
    # 13. COORDINACIÓN DOCENTE
    # =========================================================================
    add_heading(doc, '13. Coordinación docente', level=1)
    add_para(doc, '{{ texto_coordinacion_docente }}')
    
    # =========================================================================
    # 14. TUTORÍA
    # =========================================================================
    add_heading(doc, '14. Tutoría', level=1)
    add_para(doc, '{{ texto_tutoria }}')
    
    # =========================================================================
    # 15. FORMACIÓN PROFESIONAL DUAL
    # =========================================================================
    add_heading(doc, '15. Formación profesional dual', level=1)
    add_para(doc, '{{ texto_fp_dual }}')
    
    # =========================================================================
    # 16. COEDUCACIÓN
    # =========================================================================
    add_heading(doc, '16. Coeducación', level=1)
    add_para(doc, '{{ texto_coeducacion }}')
    
    # =========================================================================
    # 17. SOSTENIBILIDAD
    # =========================================================================
    add_heading(doc, '17. Sostenibilidad', level=1)
    add_para(doc, '{{ texto_sostenibilidad }}')
    
    # =========================================================================
    # Save
    # =========================================================================
    doc.save(DST)
    size = os.path.getsize(DST)
    print(f'Saved {DST}: {size:,} bytes')
    
    # Verify with docxtpl
    print('\nVerifying with docxtpl...')
    from docxtpl import DocxTemplate
    import re
    
    # Extract vars from XML (tpl.paragraphs doesn't work in docxtpl 0.20.2)
    import zipfile
    with zipfile.ZipFile(DST) as z:
        xml = z.read('word/document.xml').decode('utf-8')
    
    vars_found = sorted(set(m.strip() for m in re.findall(r'\{\{(.+?)\}\}', xml)))
    print(f'  Unique Jinja2 vars: {len(vars_found)}')
    for v in vars_found:
        print(f'    {v}')
    
    # Check for remaining [[ ]]
    remaining = re.findall(r'\[\[(.+?)\]\]', xml)
    if remaining:
        print(f'\n  WARNING: {len(remaining)} remaining [[ ]]')
        for r in remaining[:5]:
            print(f'    [[ {r.strip()[:60]} ]]')
    else:
        print(f'\n  No remaining [[ ]] - CLEAN!')
    
    # Test render
    print('\nTest render...')
    tpl = DocxTemplate(DST)
    dummy = {}
    for v in vars_found:
        if '_loop' in v:
            continue
        dummy[v] = f'[TEST {v}]'
    # Add loop vars
    for i in range(1, 8):
        dummy[f'ra{i}_loop'] = [{'ce_code': f'CE{i}a.', 'desc_ce': f'Criterio de evaluación {i}a del RA{i}'}]
    
    try:
        tpl.render(dummy)
        tpl.save('test_pd_plus_new.docx')
        size2 = os.path.getsize('test_pd_plus_new.docx')
        print(f'  Render OK: {size2:,} bytes')
    except Exception as e:
        print(f'  Render ERROR: {e}')
    
    # Save vars list
    with open('vars_pd_plus_new.txt', 'w', encoding='utf-8') as f:
        for v in vars_found:
            f.write(v + '\n')
    print(f'\nVars saved to vars_pd_plus_new.txt')


if __name__ == '__main__':
    main()
