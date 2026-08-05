# -*- coding: utf-8 -*-
"""
pdf_parte_incidencias.py
PDF A4 vertical - Parte de incidencias / justificante de falta: un
justificante de una ausencia o incidencia concreta, para que la familia lo
firme. Documento de relleno manual (fecha y motivo los indica el
profesorado al generarlo; el resto se firma a mano).
"""
import io
import pandas as pd
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, portrait
from reportlab.platypus import BaseDocTemplate, Frame, PageTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_LEFT, TA_CENTER


def _draw_page_decorations(canv, doc):
    canv.saveState()
    W, H = portrait(A4)
    canv.setFont("Helvetica-Bold", 10)
    canv.setFillColor(colors.HexColor("#777777"))
    canv.drawCentredString(W / 2, H - 1.5 * cm, doc.cal_titulo)
    canv.setFont("Helvetica", 9)
    canv.drawRightString(W - 1 * cm, 1 * cm, doc.cal_pie)
    canv.restoreState()


def _nombre_alumno(al_id, df_al):
    if df_al.empty or "ID" not in df_al.columns or al_id not in df_al["ID"].values:
        return al_id
    al = df_al[df_al["ID"] == al_id].iloc[0]
    apellidos = str(al.get("Apellidos", "") or "")
    nombre = str(al.get("Nombre", "") or "")
    return f"{apellidos}, {nombre}" if (apellidos or nombre) else al_id


def generar_pdf_parte_incidencia(info_modulo, al_id, df_al, fecha_incidencia=None, motivo_incidencia=None):
    buffer = io.BytesIO()
    W, H = portrait(A4)
    left_m, right_m, top_m, bottom_m = 2.5 * cm, 2.5 * cm, 2.0 * cm, 1.5 * cm

    doc = BaseDocTemplate(buffer, pagesize=portrait(A4), leftMargin=left_m, rightMargin=right_m,
                           topMargin=top_m, bottomMargin=bottom_m)
    doc.cal_titulo = f"Parte de incidencias  ·  {info_modulo.get('modulo', 'Módulo')}"
    doc.cal_pie = f"{info_modulo.get('centro', '')} ({info_modulo.get('profesorado', '')})"
    frame = Frame(left_m, bottom_m, W - left_m - right_m, H - top_m - bottom_m, id="main")
    doc.addPageTemplates([PageTemplate(id="port", frames=[frame], onPage=_draw_page_decorations)])

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=16, alignment=TA_CENTER, spaceAfter=20)
    norm = ParagraphStyle("Nor", parent=styles["Normal"], fontSize=11, leading=16)
    label = ParagraphStyle("Lbl", parent=styles["Normal"], fontSize=10, leading=14, fontName="Helvetica-Bold")

    nombre = _nombre_alumno(al_id, df_al)
    elements = [
        Paragraph("Justificante de falta / parte de incidencias", h1),
        Paragraph(f"<b>Alumno/a:</b> {nombre} ({al_id})", norm),
        Paragraph(f"<b>Módulo:</b> {info_modulo.get('modulo', '')}", norm),
        Paragraph(f"<b>Fecha de la incidencia:</b> {fecha_incidencia or '_______________'}", norm),
        Spacer(1, 10),
        Paragraph("Motivo:", label),
    ]
    motivo_box = Table([[Paragraph(motivo_incidencia or "", norm)]], colWidths=[W - left_m - right_m], rowHeights=[3 * cm])
    motivo_box.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#999999")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 8), ("LEFTPADDING", (0, 0), (-1, -1), 8),
    ]))
    elements.append(motivo_box)
    elements.append(Spacer(1, 40))

    firma_style = ParagraphStyle("Firma", parent=styles["Normal"], fontSize=10, alignment=TA_CENTER)
    firma_tbl = Table([
        [Paragraph("_______________________________", firma_style), Paragraph("_______________________________", firma_style)],
        [Paragraph("Fdo.: Familia / tutor/a legal", firma_style), Paragraph("Fdo.: El/la profesor/a del módulo", firma_style)],
    ], colWidths=[(W - left_m - right_m) / 2] * 2)
    elements.append(firma_tbl)

    doc.build(elements)
    buffer.seek(0)
    return buffer


def generar_docx_parte_incidencia(info_modulo, al_id, df_al, fecha_incidencia=None, motivo_incidencia=None):
    from docx_helpers import new_document, add_title, doc_to_bytes
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    nombre = _nombre_alumno(al_id, df_al)
    doc = new_document(landscape=False)
    add_title(doc, "Justificante de falta / parte de incidencias")

    doc.add_paragraph(f"Alumno/a: {nombre} ({al_id})")
    doc.add_paragraph(f"Módulo: {info_modulo.get('modulo', '')}")
    doc.add_paragraph(f"Fecha de la incidencia: {fecha_incidencia or '_______________'}")
    doc.add_heading("Motivo", level=2)
    doc.add_paragraph(motivo_incidencia or "")
    doc.add_paragraph()
    doc.add_paragraph()

    firma = doc.add_table(rows=2, cols=2)
    firma.cell(0, 0).text = "_______________________________"
    firma.cell(0, 1).text = "_______________________________"
    firma.cell(1, 0).text = "Fdo.: Familia / tutor/a legal"
    firma.cell(1, 1).text = "Fdo.: El/la profesor/a del módulo"
    for row in firma.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc_to_bytes(doc)
