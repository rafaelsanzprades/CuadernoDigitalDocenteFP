import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
try:
    from docx2pdf import convert
except ImportError:
    convert = None

def _add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h

def _set_cell_background(cell, color_hex):
    # Set background color of a cell using oxml
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def generate(data, out_docx, out_pdf):
    doc = Document()
    
    # --- Ajustes globales para comprimir en 1 página ---
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(9)
    style_normal.paragraph_format.space_before = Pt(0)
    style_normal.paragraph_format.space_after = Pt(2)
    style_normal.paragraph_format.line_spacing = 1.0
    
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Arial'
    style_h1.font.size = Pt(11)
    style_h1.paragraph_format.space_before = Pt(4)
    style_h1.paragraph_format.space_after = Pt(2)
    
    try:
        style_list = doc.styles['List Bullet']
        style_list.font.name = 'Arial'
        style_list.font.size = Pt(9)
        style_list.paragraph_format.space_before = Pt(0)
        style_list.paragraph_format.space_after = Pt(0)
    except KeyError:
        pass
    
    # --- Portada / Encabezado ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Resumen de la Programación didáctica para el alumnado")
    run.font.size = Pt(14)
    run.font.bold = True
    
    p_entradilla = doc.add_paragraph()
    p_entradilla.add_run("Seguro que tienes muchas preguntas, ¡genial!, en esta guía resumen se ha pretendido darles respuesta.").italic = True
    p_entradilla.paragraph_format.space_after = Pt(6)

    # --- H1 1 ---
    _add_heading(doc, "¿Cómo contribuirá este módulo en mi Perfil profesional?", 1)
    doc.add_paragraph("Según indica el art. 7 apartado 1, Orden ECD/884/2016, 15 julio 2016, ejerce su actividad por cuenta ajena en empresas de montaje y mantenimiento de instalaciones electrotécnicas de edificios, viviendas, oficinas, locales comerciales e industriales. La formación de este módulo se relaciona con los Objetivos generales (OG) del ciclo (art. 9): a), b), c), d), e), f), g) y h), r), s), t), u), v), w) y x); y las Competencias Profesionales, Personales y Sociales (CPPS) (art. 5): a), b), c), d), e), f), g*), h), p), q), r), s), t), u) y v).")

    # --- H1 2 ---
    _add_heading(doc, "¿En qué seré competente?", 1)
    doc.add_paragraph("La competencia general de este título consiste, como expresa el art. 4, Orden ECD/884/2016, 15 julio 2016, en realizar operaciones auxiliares en el montaje y mantenimiento de elementos y equipos eléctricos y electrónicos, así como en instalaciones electrotécnicas y de telecomunicaciones para edificios y conjuntos de edificios, aplicando las técnicas requeridas, operando con la calidad indicada, observando las normas de prevención de riesgos laborales y protección medioambiental.")

    # --- H1 3 ---
    _add_heading(doc, "¿Qué sabré hacer? Se le denomina: Resultados de Aprendizaje", 1)
    df_ra = data.get("df_ra", [])
    if df_ra:
        for ra in df_ra:
            id_str = str(ra.get('id_ra', ''))
            prefix = "" if id_str.upper().startswith("RA") else "RA"
            doc.add_paragraph(f"{prefix}{id_str}. {ra.get('desc_ra', '')}", style='List Bullet')
    else:
        doc.add_paragraph("RA1. Selecciona los elementos, equipos y herramientas para la realización del montaje y mantenimiento.", style='List Bullet')
        doc.add_paragraph("RA2. Monta canalizaciones, soportes y cajas en baja tensión y/o domóticas.", style='List Bullet')
        doc.add_paragraph("RA3. Tiende el cableado entre equipos y elementos de baja tensión y/o domóticas.", style='List Bullet')
        doc.add_paragraph("RA4. Instala mecanismos y elementos identificando sus componentes y aplicaciones.", style='List Bullet')
        doc.add_paragraph("RA5. Realiza operaciones auxiliares de mantenimiento.", style='List Bullet')

    # --- H1 4 ---
    _add_heading(doc, "¿Qué aprenderé? Son los Contenidos mínimos", 1)
    df_ud = data.get("df_ud", [])
    if df_ud:
        for ud in df_ud:
            id_str = str(ud.get('id_ud', ''))
            prefix = "" if id_str.upper().startswith("UD") else "UD"
            doc.add_paragraph(f"{prefix}{id_str}. {ud.get('desc_ud', '')}", style='List Bullet')
    else:
        doc.add_paragraph("UD1. Selección de elementos, equipos y herramientas de instalaciones eléctricas / domóticas.", style='List Bullet')
        doc.add_paragraph("UD2. Montaje de canalizaciones, soportes y cajas en instalaciones eléctricas de baja tensión y/o domótica.", style='List Bullet')
        doc.add_paragraph("UD3. Tendido de cableado entre equipos y elementos de instalaciones eléctricas/domóticas.", style='List Bullet')
        doc.add_paragraph("UD4. Instalación de mecanismos y elementos de las instalaciones eléctricas/domóticas.", style='List Bullet')
        doc.add_paragraph("UD5. Mantenimiento de instalaciones eléctricas y/o domóticas de edificios.", style='List Bullet')

    # --- H1 5 ---
    _add_heading(doc, "¿Cómo aprobaré el módulo? Criterios de evaluación y calificación del módulo", 1)
    
    p1 = doc.add_paragraph()
    p1.add_run("55 % Desarrollo de las prácticas en el Aula Taller de instalaciones eléctricas.\n").bold = True
    p1.add_run("Rúbrica específica de instalaciones. Autoevaluación previa hasta tres intentos para mejorar la nota.")

    p2 = doc.add_paragraph()
    p2.add_run("10 % Corrección del Cuaderno del Taller.\n").bold = True
    p2.add_run("Apuntes de clase, resumen de cada Unidad didáctica, informes de las prácticas y anotaciones.")

    p3 = doc.add_paragraph()
    p3.add_run("5 % Preparación del examen teórico. Debate en grupo.\n").bold = True
    p3.add_run("Ronda de preguntas verbales, nivel de participación, resolución de dudas.")

    p4 = doc.add_paragraph()
    p4.add_run("30 % Examen teórico escrito (con una calificación mínima de 5 para media).\n").bold = True
    p4.add_run("Se pregunta sobre cuestiones de aplicación sobre el contenido teórico.\n+ 1 punto adicional por actitud y comportamiento positivo.\nMenos del 5 % de faltas de asistencia a clase y sin partes negativos de comportamiento")

    # --- H1 6 ---
    _add_heading(doc, "Recordad:", 1)
    p_rec = doc.add_paragraph()
    p_rec.add_run("Más del 15 % de faltas de asistencia a clase implica la Perdida del derecho a evaluación continua.").bold = True

    p_final = doc.add_paragraph("Tu situación es similar a la del resto del alumnado que ha obtenido su titulación. ¡ANIMO!")
    p_final.paragraph_format.space_before = Pt(6)
    
    # Save DOCX
    doc.save(out_docx)
    
    # Convert to PDF
    try:
        abs_docx = os.path.abspath(out_docx)
        abs_pdf = os.path.abspath(out_pdf)
        if convert:
            convert(abs_docx, abs_pdf)
        else:
            print("docx2pdf not available, skipping PDF conversion.")
    except Exception as e:
        print(f"Error converting to PDF: {e}")
        if os.path.exists(out_pdf):
            os.remove(out_pdf)
