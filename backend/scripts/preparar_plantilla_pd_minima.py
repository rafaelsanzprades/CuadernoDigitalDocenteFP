"""
Script one-shot: genera la plantilla DOCX para PD- (resumen alumnado).

Crea un DOCX con formato compacto (1 página) y marcadores Jinja2.

Uso:
    cd backend
    python scripts/preparar_plantilla_pd_minima.py

Resultado:
    backend/templates/modelo_pd_fp-.docx
"""

import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DST_DIR = os.path.join(BASE, "templates")
DST = os.path.join(DST_DIR, "modelo_pd_fp-.docx")


def main():
    os.makedirs(DST_DIR, exist_ok=True)
    doc = Document()

    # --- Ajustes globales para comprimir en 1 página ---
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.bottom_margin = Cm(1.0)
        section.right_margin = Cm(1.0)

    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(9)
    style_normal.font.color.rgb = RGBColor(0, 0, 0)
    style_normal.paragraph_format.space_before = Pt(0)
    style_normal.paragraph_format.space_after = Pt(2)
    style_normal.paragraph_format.line_spacing = 1.0

    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Arial'
    style_h1.font.size = Pt(11)
    style_h1.font.color.rgb = RGBColor(0, 0, 0)
    style_h1.paragraph_format.space_before = Pt(4)
    style_h1.paragraph_format.space_after = Pt(2)

    try:
        style_list = doc.styles['List Bullet']
        style_list.font.name = 'Arial'
        style_list.font.size = Pt(9)
        style_list.font.color.rgb = RGBColor(0, 0, 0)
        style_list.paragraph_format.space_before = Pt(0)
        style_list.paragraph_format.space_after = Pt(0)
    except KeyError:
        pass

    # --- Portada / Encabezado ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Resumen de la Programación didáctica para el alumnado\n")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0, 0, 0)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Módulo: {{ modulo }}\nTítulo: {{ ciclo }}. Curso: {{ curso_academico }}")
    run_sub.font.name = 'Arial'
    run_sub.font.color.rgb = RGBColor(0, 0, 0)
    p_sub.paragraph_format.space_after = Pt(12)

    p_entradilla = doc.add_paragraph()
    run_ent = p_entradilla.add_run("Seguro que tienes muchas preguntas, ¡genial!, en esta guía resumen se ha pretendido darles respuesta.")
    run_ent.italic = True
    run_ent.font.name = 'Arial'
    p_entradilla.paragraph_format.space_after = Pt(6)

    # --- H1 1: Perfil profesional ---
    doc.add_heading("¿Cómo contribuirá este módulo en mi Perfil profesional?", level=1)
    doc.add_paragraph("{{ texto_perfil_profesional }}")

    # --- H1 2: Competencia ---
    doc.add_heading("¿En qué seré competente?", level=1)
    doc.add_paragraph("{{ texto_competencia_general }}")

    # --- H1 3: Resultados de Aprendizaje ---
    doc.add_heading("¿Qué sabré hacer? Se le denomina: Resultados de Aprendizaje", level=1)
    doc.add_paragraph("{%p for ra in list_ras %}")
    p_ra = doc.add_paragraph(style='List Bullet')
    p_ra.text = "{{ ra }}"
    doc.add_paragraph("{%p endfor %}")

    # --- H1 4: Contenidos (UDs) ---
    doc.add_heading("¿Qué aprenderé? Son los Contenidos mínimos", level=1)
    doc.add_paragraph("{%p for ud in list_uds %}")
    p_ud = doc.add_paragraph(style='List Bullet')
    p_ud.text = "{{ ud }}"
    doc.add_paragraph("{%p endfor %}")

    # --- H1 5: Criterios de calificación ---
    doc.add_heading("¿Cómo aprobaré el módulo? Criterios de evaluación y calificación del módulo", level=1)
    
    p_placeholder = doc.add_paragraph("[[TABLA_INSTRUMENTOS]]")
    p_placeholder.runs[0].font.name = 'Arial'
    
    doc.add_paragraph() # Add a small spacing after the table

    # --- H1 6: Recordad ---
    doc.add_heading("Recordad:", level=1)
    p_rec = doc.add_paragraph()
    r_rec = p_rec.add_run("{{ texto_recordatorio }}")
    r_rec.bold = True
    r_rec.font.name = 'Arial'

    p_final = doc.add_paragraph("{{ texto_final }}")
    p_final.runs[0].font.name = 'Arial'
    p_final.paragraph_format.space_before = Pt(6)

    # --- Guardar ---
    doc.save(DST)
    print(f"[OK] Plantilla guardada en: {DST}")


if __name__ == "__main__":
    main()
