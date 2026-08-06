"""
Script one-shot: corrige defectos de autoría heredados en la plantilla DOCX
de la PD= (Aragón/BOA), detectados al generar el documento real con datos
demo y compararlo con lo que el profesor esperaba ver.

Corrige, en `modelo_pd_fp=.docx`:
  1. Quita todos los resaltados en amarillo (marcas de cuando se adaptó la
     plantilla oficial del BOA: título, aviso de actualización normativa,
     campo de índice, cabecera de la sección A).
  2. Sección H (RESULTADOS DE APRENDIZAJE): los placeholders usaban por
     error ra8_texto_c2..ra14_texto_c2 (nunca definidos por el generador,
     de ahí que saliera en blanco) en vez de ra1_texto_c2..ra7_texto_c2
     (que ya están bien rellenos, son los mismos textos de la sección C2).
     También elimina un bucle Jinja huérfano ({%p for ra in list_ras %})
     que quedó a medio cablear y nunca se usa.
  3. Cabecera de página: en la fila CURSO ACADÉMICO el valor solo ocupaba
     una columna de la tabla y quedaba una celda vacía suelta antes de la
     celda de página (a diferencia de CICLO/MÓDULO, cuyo valor ya fusiona
     3 columnas) — se fusiona para seguir el mismo patrón. Además solo
     había campo PAGE; se añade " de " + NUMPAGES para que quede
     "Página: X de Y".

Uso:
    cd backend
    python scripts/arreglar_plantilla_pd_suficiente.py

Resultado:
    Sobrescribe backend/templates/modelo_pd_fp=.docx
"""

import copy
import os
import re

from docx import Document
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_PATH = os.path.join(BASE, "templates", "modelo_pd_fp=.docx")


def _quitar_resaltado_amarillo_parrafos(paragraphs):
    for p in paragraphs:
        for r in p.runs:
            r.font.highlight_color = None


def quitar_resaltado_amarillo(doc):
    _quitar_resaltado_amarillo_parrafos(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                _quitar_resaltado_amarillo_parrafos(cell.paragraphs)
    for section in doc.sections:
        _quitar_resaltado_amarillo_parrafos(section.header.paragraphs)
        for t in section.header.tables:
            for row in t.rows:
                for cell in row.cells:
                    _quitar_resaltado_amarillo_parrafos(cell.paragraphs)


def _set_paragraph_text(paragraph, text):
    """Reescribe el texto de un párrafo en un único run, conservando el
    formato del primer run — el texto suele venir partido en varios runs
    por los saltos de autocorrección de Word, así que un replace ingenuo
    sobre el texto concatenado no basta."""
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for extra in paragraph.runs[1:]:
        extra.text = ""


def arreglar_seccion_h(doc):
    patron = re.compile(r"^RA (\d+)\. \{\{ ra(\d+)_texto_c2 \}\}$")
    paras = doc.paragraphs
    for p in paras:
        text = p.text.strip()
        m = patron.match(text)
        if m and int(m.group(2)) >= 8:
            n = m.group(1)
            _set_paragraph_text(p, f"RA {n}. {{{{ ra{n}_texto_c2 }}}}")

    # El bucle huérfano "{%p for ra in list_ras %}{{ ra }}{%p endfor %}" son
    # 3 párrafos consecutivos y exactos — se localizan por posición (a partir
    # del marcador "for ra in list_ras", único en el documento) en vez de por
    # texto genérico "{%p endfor %}", que también usa el bucle (no huérfano,
    # no tocar) de "{%p for ud in list_uds %}" un poco más arriba.
    idx_inicio = next(
        (i for i, p in enumerate(paras) if "for ra in list_ras" in p.text),
        None,
    )
    if idx_inicio is not None:
        bloque = paras[idx_inicio:idx_inicio + 3]
        textos = [p.text.strip() for p in bloque]
        assert textos == ["{%p for ra in list_ras %}", "{{ ra }}", "{%p endfor %}"], textos
        for p in bloque:
            p._p.getparent().remove(p._p)


def fusionar_columnas_cabecera(doc):
    table = doc.sections[0].header.tables[0]
    row = table.rows[2]  # fila CURSO ACADÉMICO
    row.cells[2].merge(row.cells[3])


def anadir_paginas_totales(doc):
    table = doc.sections[0].header.tables[0]
    row = table.rows[2]
    cell_pagina = row.cells[4]
    p = cell_pagina.paragraphs[0]
    runs = p.runs
    # runs: [0]="Página: ", [1]=begin PAGE, [2]=instrText PAGE,
    #       [3]=separate, [4]=valor cacheado, [5]=end
    if len(runs) < 6:
        return

    de_run = copy.deepcopy(runs[0]._r)
    for t in de_run.findall(qn("w:t")):
        t.text = " de "
        t.set(qn("xml:space"), "preserve")

    new_begin = copy.deepcopy(runs[1]._r)
    new_instr = copy.deepcopy(runs[2]._r)
    for t in new_instr.findall(qn("w:instrText")):
        t.text = "NUMPAGES"
    new_sep = copy.deepcopy(runs[3]._r)
    new_val = copy.deepcopy(runs[4]._r)
    for t in new_val.findall(qn("w:t")):
        t.text = "1"
    new_end = copy.deepcopy(runs[5]._r)

    prev = runs[5]._r
    for elem in (de_run, new_begin, new_instr, new_sep, new_val, new_end):
        prev.addnext(elem)
        prev = elem


def main():
    doc = Document(TEMPLATE_PATH)
    quitar_resaltado_amarillo(doc)
    arreglar_seccion_h(doc)
    fusionar_columnas_cabecera(doc)
    anadir_paginas_totales(doc)
    doc.save(TEMPLATE_PATH)
    print(f"[OK] Plantilla corregida y guardada en: {TEMPLATE_PATH}")


if __name__ == "__main__":
    main()
