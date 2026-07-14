"""
Script para convertir la plantilla PD+ (JEG) a Jinja2.

Abre el .docx como ZIP, modifica el XML word/document.xml directamente.
V2: Mapeo completo con limpieza de bloques explicativos y campos unicode.

Lee: backend/templates/modelo_pd_fp+.docx
Escribe: backend/templates/modelo_pd_fp+_tpl.docx
"""

import os
import re
import sys
import zipfile
from lxml import etree

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BACKEND_DIR = os.path.dirname(SCRIPT_DIR)
TEMPLATE_INPUT = os.path.join(BACKEND_DIR, 'templates', 'modelo_pd_fp+.docx')
TEMPLATE_OUTPUT = os.path.join(BACKEND_DIR, 'templates', 'modelo_pd_fp+_tpl.docx')

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
XML_NS = 'http://www.w3.org/XML/1998/namespace'


def get_full_text(para):
    texts = []
    for t in para.iter(f'{{{W_NS}}}t'):
        if t.text:
            texts.append(t.text)
    return ''.join(texts)


def replace_in_paragraph(para, old_text, new_text):
    full = get_full_text(para)
    if old_text not in full:
        return False
    new_full = full.replace(old_text, new_text)
    t_elems = list(para.iter(f'{{{W_NS}}}t'))
    if not t_elems:
        return False
    t_elems[0].text = new_full
    t_elems[0].set(f'{{{XML_NS}}}space', 'preserve')
    for t in t_elems[1:]:
        t.text = ''
        if f'{{{XML_NS}}}space' in t.attrib:
            del t.attrib[f'{{{XML_NS}}}space']
    return True


def delete_paragraph(para):
    """Marca un parrafo para eliminacion vaciandolo."""
    for t in para.iter(f'{{{W_NS}}}t'):
        t.text = ''
    return True


def process_xml(xml_bytes):
    root = etree.fromstring(xml_bytes)

    # === MAPEO COMPLETO (patron exacto en texto del parrafo -> reemplazo) ===
    replacements = [
        # --- Portada ---
        ('[[ Módulo profesional | Proyecto ]]', '{{ tipo_elemento }}'),
        ('[[ módulo profesional | proyecto ]]', '{{ tipo_elemento }}'),
        ('[[ módulo | proyecto ]]', '{{ tipo_elemento }}'),
        ('[[ D | E ]]', '{{ grado }}'),
        ('[[ 1 | 2 | 3 ]]', '{{ nivel }}'),
        ('[[ 1.º | 2.º | 3.º ]]', '{{ curso_numero }}'),
        ('[[ presencial | semipresencial | virtual ]]', '{{ modalidad }}'),

        # --- Contextualizacion: centro ---
        ('[[ nombre del centro ]]', '{{ nombre_centro }}'),
        ('[[ público | concertado | privado ]]', '{{ tipo_centro }}'),
        ('[[ municipio | comarca | provincia ]]', '{{ municipio }}'),
        ('[[ tradiciones | patrimonio cultural | idioma o lenguas cooficiales | festividades relevantes ]]', '{{ contexto_cultural }}'),
        ('[[ transporte público | infraestructuras viarias | conectividad digital | condiciones de acceso ]]', '{{ contexto_accesibilidad }}'),
        ('[[ sectores productivos predominantes | actividades económicas más relevantes | industrias clave ]]', '{{ sectores_productivos }}'),
        ('[[ grado D | grado E ]]', '{{ grado }}'),
        ('[[ perfiles profesionales demandados | evolución del empleo | nichos emergentes ]]', '{{ perfiles_profesionales }}'),
        ('[[ asociaciones sectoriales | viveros de empresas | parques tecnológicos | centros de innovación ]]', '{{ entorno_empresarial }}'),
        ('[[ número de alumnado | número de profesorado | número de grupos ]]', '{{ dimension_centro }}'),
        ('[[ elementos que definen su identidad institucional | carácter propio | valores educativos relevantes ]]', '{{ identidad_centro }}'),

        # --- Contextualizacion: alumnado ---
        ('[[ franja de edad ]]', '{{ franja_edad }}'),
        ('[[ homogéneo | disperso | variado ]]', '{{ nivel_madurez }}'),
        ('[[ ESO | Bachillerato | prueba de acceso | otros ciclos formativos | otras vías ]]', '{{ vias_acceso }}'),
        ('[[ titulación de acceso ]]', '{{ titulacion_acceso }}'),
        ('[[ relacionada | no relacionada ]]', '{{ relacionada_modulo }}'),
        ('[[ alto | medio | bajo ]]', '{{ grado_implicacion }}'),
        ('[[ inserción laboral | continuidad de estudios | desarrollo personal | otros objetivos ]]', '{{ expectativas }}'),

        # --- Desarrollo curricular ---
        ('[[ horas ]]', '{{ horas }}'),
        ('[[ curso escolar ]]', '{{ curso_escolar }}'),
        ('[[ curso académico ]]', '{{ curso_academico }}'),
        ('[[ código del grado ]]', '{{ codigo_grado }}'),
        ('[[ denominación del grado ]]', '{{ denominacion_grado }}'),

        # --- Organizacion ---
        ('[[ secuencial | paralela | mixta ]]', '{{ organizacion_secuencial }}'),

        # --- FEOE ---
        ('[[ general | intensivo ]]', '{{ regimen_feoe }}'),
        ('[[ fecha de inicio ]]', '{{ fecha_inicio }}'),
        ('[[ fecha de fin ]]', '{{ fecha_fin }}'),
        ('[[ número de horas ]]', '{{ horas_feoe }}'),
        ('[[ n.º RA en FEOE ]]', '{{ num_ra_feoe }}'),
        ('[[ n.º total de RA ]]', '{{ num_total_ra }}'),
        ('[[ descripción de la organización ]]', '{{ organizacion_no_feoe }}'),

        # --- Metodologia ---
        ('[[ específica | polivalente | informática ]]', '{{ tipo_aula }}'),
        ('[[ de sobremesa | portátiles ]]', '{{ tipo_equipos }}'),
        ('[[ Aeducar ]]', '{{ plataforma_educativa }}'),
        ('[[ nombre, versión ]]', '{{ software_nombre }}'),
        ('[[ Otros ]]', '{{ otros_recursos }}'),
        ('[[ vídeos formativos, tutoriales… ]]', '{{ recursos_multimedia }}'),

        # --- Agrupamientos ---
        ('[[ número ]]', '{{ tamanno_equipo }}'),

        # --- Desdobles ---
        ('[[ proyecto | reto | proyecto/reto ]]', '{{ tipo_colaborativo }}'),

        # --- Evaluacion ---
        ('[[ calificación ]]', '{{ calificacion_minima }}'),

        # --- Actividades complementarias ---
        ('[[ complementaria | extraescolar ]]', '{{ tipo_actividad }}'),
        ('[[ Descripción de la actividad ]]', '{{ actividad_descripcion }}'),
        ('[[ Descripción de la actividad de evaluación y/o calificación, si procede ]]', '{{ actividad_evaluacion }}'),

        # --- Recuperacion ---
        ('[[ correo electrónico | la plataforma educativa Aeducar |  ... ]]', '{{ medio_entrega }}'),
        ('[[ correo electrónico | la plataforma educativa Aeducar | presencialmente | ... ]]', '{{ medio_evaluacion }}'),
        ('[[ correo electrónico |  la plataforma educativa Aeducar, en foros y por mensajería interna ]]', '{{ medio_dudas }}'),
        ('[[ en una reunión presencial | mediante una comunicación escrita | ... ]]', '{{ formato_comunicacion_inicial }}'),
        ('[[ por correo electrónico | mediante la plataforma Aeducar | … ]]', '{{ formato_comunicacion_desarrollo }}'),

        # --- Contingencia ---
        ('[[ correo electrónico | … ]]', '{{ medio_comunicacion }}'),
        ('[[ llamadas telefónicas | reuniones presenciales | … ]]', '{{ medio_seguimiento }}'),
        ('[[ unidades didácticas | resultados de aprendizaje | ... ]]', '{{ ambito_recuperacion }}'),
        ('[[ un repositorio del departamento | curso correspondiente de la plataforma Aeducar | carpeta compartida ]]', '{{ repositorio_recursos }}'),
        ('[[ la coordinación del departamento didáctico correspondiente | supervisión de la persona tutora ]]', '{{ coordinacion_contingencia }}'),
        ('[[ el curso correspondiente de la plataforma Aeducar | ... ]]', '{{ ubicacion_recursos }}'),
        ('[[ el profesorado responsable del módulo profesional o proyecto ... ]]', '{{ profesorado_contingencia }}'),

        # --- Contextualizacion: campos largos restantes ---
        ('[[ tamaño de la población | distribución por edades | densidad poblacional | tendencia demográfica: crecimiento, estabilidad o decrecimiento ]]', '{{ contexto_demografico }}'),
        ('[[ ESO | Bachillerato | ciclos formativos | cursos de especialización | programas de atención a la diversidad | otros ]]', '{{ oferta_formativa }}'),
        ('[[ programas de innovación | planes de mejora | proyectos Erasmus+ | iniciativas de sostenibilidad | planes de digitalización ]]', '{{ programas_centro }}'),
        ('[[ grado de autonomía | hábitos de estudio | capacidad de trabajo en equipo | estilos y ritmos de aprendizaje | clima y dinámica de grupo ]]', '{{ caracteristicas_alumnado_extra }}'),

        # --- Evaluacion parciales ---
        ('[[ tres evaluaciones parciales, una al finalizar cada trimestre | dos evaluaciones parciales, una por cada evaluación ]]', '{{ num_evaluaciones_parciales }}'),

        # --- Recuperacion pendientes ---
        ('[[ el inicio de las actividades lectivas | fecha de inicio | … ]]', '{{ fecha_inicio_plan_recuperacion }}'),
        ('[[ la primera evaluación final | fecha de finalización | … ]]', '{{ fecha_fin_plan_recuperacion }}'),

        # --- Contingencia: campos restantes ---
        ('[[ un repositorio del departamento | curso correspondiente de la plataforma Aeducar | carpeta compartida ]]', '{{ repositorio_recursos }}'),
        ('[[ la coordinación del departamento didáctico correspondiente | supervisión de la persona tutora del grupo ]]', '{{ coordinacion_contingencia }}'),
        ('[[ evaluará la situación del grupo | adaptará la temporización, si fuera necesario | ... ]]', '{{ actuaciones_contingencia }}'),
        ('[[ el curso correspondiente de la plataforma Aeducar | en formato impreso | ... ]]', '{{ formato_comunicacion }}'),

        # --- Genericos (puntos suspensivos unicode) ---
        ('[[ … ]]', '{{ campo_libre }}'),
        ('[[ …]]', '{{ campo_libre }}'),
        ('[[...]]', '{{ campo_libre }}'),
        ('[[ ... ]]', '{{ campo_libre }}'),

        # --- Tablas de evaluacion ---
        ('[[ 1.ª evaluación | 2.ª evaluación | 3.ª evaluación | … ]]', '{{ actividad_temporizacion }}'),
    ]

    # === BLOQUES EXPLICATIVOS A ELIMINAR ===
    delete_patterns = [
        'OPCIÓN A: MÓDULOS PROFESIONALES QUE DESARROLLAN ALGÚN RA EN LA EMPRESA U ORGANISMO EQUIPARADO.',
        'OPCIÓN B: MÓDULOS PROFESIONALES QUE NO DESARROLLAN NINGÚN RA EN LA EMPRESA U ORGANISMO EQUIPARADO.',
        'Este módulo profesional no está asociado a ningún estándar de competencia profesional.',
        'Este {{ tipo_elemento }} no contempla la aplicación de desdobles.',
        'modificación pendiente de publicación',
        'La evaluación y calificación del alumnado se ajustará a lo indicado en el apartado 4.7',
        'El título o el currículo del grado correspondiente al que pertenece este',
    ]

    count = 0
    deleted = 0
    for para in root.iter(f'{{{W_NS}}}p'):
        text = get_full_text(para)

        # Verificar si es un bloque explicativo a eliminar
        should_delete = False
        for dp in delete_patterns:
            if dp in text:
                should_delete = True
                break

        # Tambien eliminar bloques explicativos largos con [[ ... ]]
        if '[[' in text and len(text) > 300:
            should_delete = True

        if should_delete:
            delete_paragraph(para)
            deleted += 1
            continue

        if '[[' not in text:
            continue

        for old, new in replacements:
            if old in text:
                if replace_in_paragraph(para, old, new):
                    count += 1
                    # Re-leer texto por si hay multiples campos
                    text = get_full_text(para)

    # Procesar tablas
    for tbl in root.iter(f'{{{W_NS}}}tbl'):
        for para in tbl.iter(f'{{{W_NS}}}p'):
            text = get_full_text(para)
            if '[[' not in text:
                continue
            for old, new in replacements:
                if old in text:
                    if replace_in_paragraph(para, old, new):
                        count += 1
                        text = get_full_text(para)

    return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True), count, deleted


def convert_ud_ra_tables(xml_bytes):
    """
    Tercer pase: convierte textos estáticos de tablas de UD/RA a variables Jinja2.
    
    - "Título UD N" → "{{ udN_titulo }}"
    - "Título del módulo profesional o proyecto" en celdas de UD → "{{ udN_titulo }}"
    - Cabeceras de tabla RA×CE se mantienen (son estructura fija)
    """
    root = etree.fromstring(xml_bytes)
    count = 0
    
    # Patrón para títulos de UD
    ud_patterns = [
        (re.compile(r'T[ií]tulo UD\s*(\d+)'), lambda m: f'{{{{ ud{m.group(1)}_titulo }}}}'),
        (re.compile(r'T[ií]tulo del m[oó]dulo profesional o proyecto'), lambda m: '{{ modulo }}'),
    ]
    
    # Procesar todas las tablas
    for tbl in root.iter(f'{{{W_NS}}}tbl'):
        for para in tbl.iter(f'{{{W_NS}}}p'):
            text = get_full_text(para)
            if not text or not text.strip():
                continue
            
            for pattern, replacement in ud_patterns:
                new_text = pattern.sub(replacement, text)
                if new_text != text:
                    t_elems = list(para.iter(f'{{{W_NS}}}t'))
                    if t_elems:
                        t_elems[0].text = new_text
                        t_elems[0].set(f'{{{XML_NS}}}space', 'preserve')
                        for t in t_elems[1:]:
                            t.text = ''
                            if f'{{{XML_NS}}}space' in t.attrib:
                                del t.attrib[f'{{{XML_NS}}}space']
                        count += 1
                        text = new_text
    
    return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True), count


def cleanup_remaining(xml_bytes):
    """
    Segundo pase: convierte cualquier [[ ... ]] restante a {{ campo_libre }}
    y elimina campos anidados tipo [[ texto [[ subcampo ]].
    """
    root = etree.fromstring(xml_bytes)
    count = 0

    # Regex para [[ ... ]] con cualquier contenido (incluyendo unicode …)
    bracket_re = re.compile(r'\[\[.*?\]\]')

    for para in list(root.iter(f'{{{W_NS}}}p')) + list(root.iter(f'{{{W_NS}}}tbl')):
        for p in para.iter(f'{{{W_NS}}}p') if para.tag == f'{{{W_NS}}}tbl' else [para]:
            text = get_full_text(p)
            if '[[' not in text:
                continue
            new_text = bracket_re.sub('{{ campo_libre }}', text)
            if new_text != text:
                t_elems = list(p.iter(f'{{{W_NS}}}t'))
                if t_elems:
                    t_elems[0].text = new_text
                    t_elems[0].set(f'{{{XML_NS}}}space', 'preserve')
                    for t in t_elems[1:]:
                        t.text = ''
                        if f'{{{XML_NS}}}space' in t.attrib:
                            del t.attrib[f'{{{XML_NS}}}space']
                    count += 1

    return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True), count


def main():
    if not os.path.exists(TEMPLATE_INPUT):
        print(f"No se encontro la plantilla: {TEMPLATE_INPUT}")
        sys.exit(1)

    print(f"Leyendo: {TEMPLATE_INPUT}")

    tmp_output = TEMPLATE_OUTPUT + '.tmp'
    total_count = 0
    total_deleted = 0

    with zipfile.ZipFile(TEMPLATE_INPUT, 'r') as zin:
        with zipfile.ZipFile(tmp_output, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    data, count, deleted = process_xml(data)
                    total_count = count
                    total_deleted = deleted
                    # Segundo pase: convertir títulos de UD/RA a Jinja2
                    data, ud_count = convert_ud_ra_tables(data)
                    print(f"  word/document.xml: {count} reemplazos, {deleted} bloques eliminados, {ud_count} conversiones UD/RA")
                    # Tercer pase: limpiar campos [[ ]] restantes
                    data, cleanup_count = cleanup_remaining(data)
                    print(f"  word/document.xml: {cleanup_count} limpiezas adicionales")
                zout.writestr(item, data)

    if os.path.exists(TEMPLATE_OUTPUT):
        os.remove(TEMPLATE_OUTPUT)
    os.rename(tmp_output, TEMPLATE_OUTPUT)

    print(f"Guardado: {TEMPLATE_OUTPUT}")

    # Verificar
    from docx import Document
    doc = Document(TEMPLATE_OUTPUT)
    remaining = sum(1 for p in doc.paragraphs if '[[' in p.text)
    jinja_vars = set()
    for p in doc.paragraphs:
        for m in re.findall(r'\{\{\s*(\w+)\s*\}\}', p.text):
            jinja_vars.add(m)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for m in re.findall(r'\{\{\s*(\w+)\s*\}\}', p.text):
                        jinja_vars.add(m)

    print(f"Campos [[ ]] restantes: {remaining}")
    print(f"Variables Jinja2 unicas: {len(jinja_vars)}")
    for v in sorted(jinja_vars):
        print(f"  {v}")


if __name__ == '__main__':
    main()
