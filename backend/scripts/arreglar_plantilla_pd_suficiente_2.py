"""
Script one-shot (2ª pasada): convierte a dinámicas las secciones de la
plantilla PD= que estaban con un número fijo de RA/UD/CE hardcodeado
(copiado de un módulo real de 7 RA y 11 UD), a partir de las anotaciones
"RF." que Rafael dejó escritas directamente en el documento.

Corrige, en `modelo_pd_fp=.docx`:
  1. Quita la anotación de cálculo "{{ RAFA: 100%-15%PdEvc}}" que quedó
     dentro de un tag Jinja real (rompía el render con un error de sintaxis).
  2. Quita las 9 anotaciones "RF. ..." dejadas como marcador de trabajo.
  3. Las dos listas de "unidades formativas" (Sección A y "RELACIÓN DE
     UNIDADES FORMATIVAS...") usaban el tag genérico sin numerar
     "{{ ud_item }}" (nunca definido por el generador, siempre en blanco) —
     se convierten en un bucle dinámico sobre list_uf_items.
  4. Sección C1: se borra la lista fija "UD 1...UD 11" — ya existía un
     bucle dinámico sin usar justo encima ({%p for ud in list_uds %}) que
     ahora sí queda cableado con el formato que necesita esta sección.
  5. Sección C: se sustituye la tabla estática (7 columnas de RA fijas) por
     el marcador [[TABLA_SECUENCIACION]], que generador_pd_suficiente_tpl.py
     rellena con una tabla construida con python-docx con tantas columnas
     de RA y filas de UD como tenga el módulo real.
  6. Sección C2: se sustituye el bloque estático "RA 1...RA 7" con sus CE
     hardcodeados por un bucle dinámico sobre list_c2 (RA + sus CE reales,
     cualquier número de cada uno).
  7. Sección C3: se sustituyen las líneas de porcentajes fijos por el
     marcador [[TABLA_INSTRUMENTOS_C3]] (misma tabla que en PD-).
  8. Sección H: se sustituye la lista fija "RA 1...RA 7" por un bucle
     dinámico sobre list_ras_h, y el "seis" hardcodeado de "consta de seis
     Resultados de aprendizaje" por {{ n_ra }}.
  9. Sección J (Herramientas): se sustituyen las 4 líneas fijas por un
     bucle dinámico sobre list_herramientas (recursos_espacios resuelto).
  10. Sección N: el párrafo fijo sobre dependencia entre módulos se
      sustituye por {{ textos_pd_metodologia_labor_coordinada }}.

Uso:
    cd backend
    python scripts/arreglar_plantilla_pd_suficiente_2.py

Resultado:
    Sobrescribe backend/templates/modelo_pd_fp=.docx
"""

import os
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_PATH = os.path.join(BASE, "templates", "modelo_pd_fp=.docx")


def _set_paragraph_text(paragraph, text):
    """Reescribe el texto de un párrafo en un único run, conservando el
    formato del primer run — el texto suele venir partido en varios runs
    por los saltos de autocorrección de Word."""
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for extra in paragraph.runs[1:]:
        extra.text = ""


def _nuevo_parrafo_tras(paragraph, texto=""):
    """Inserta un párrafo nuevo (estilo Normal) justo después de `paragraph`."""
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    nuevo = Paragraph(new_p, paragraph._parent)
    if texto:
        nuevo.add_run(texto)
    return nuevo


def _borrar_parrafos(paragraphs):
    for p in paragraphs:
        p._p.getparent().remove(p._p)


def _quitar_anotaciones_sueltas(doc):
    """Notas de cálculo que Rafael dejó escritas dentro de tags Jinja reales
    (rompían el render con TemplateSyntaxError: unexpected '%')."""
    for p in doc.paragraphs:
        if "RAFA" in p.text and "PdEvc" in p.text:
            _set_paragraph_text(p, p.text.replace("{{ RAFA: 100%-15%PdEvc}} ", ""))
        elif "{{%PdEvc}}" in p.text:
            _set_paragraph_text(p, p.text.replace("{{%PdEvc}}", ""))


def _quitar_comentarios_rf(doc):
    # Algunas anotaciones quedaron escritas dentro de un tag Jinja real
    # ("{{ RF. AQUÍ PONER... }}"), así que no basta con mirar el principio
    # del párrafo — se busca "RF." en cualquier parte del texto.
    to_delete = [p for p in doc.paragraphs if "RF." in p.text]
    _borrar_parrafos(to_delete)


def _dinamizar_listas_ud_item(doc):
    """Convierte cada pareja de "{{ ud_item }}" consecutivos en un bucle
    {%p for x in list_uf_items %}{{ x }}{%p endfor %}. Hay 2 parejas en la
    plantilla (Sección A y "RELACIÓN DE UNIDADES FORMATIVAS...")."""
    while True:
        paras = doc.paragraphs
        idx = next((i for i, p in enumerate(paras) if p.text.strip() == "{{ ud_item }}"), None)
        if idx is None:
            break
        primero, segundo = paras[idx], paras[idx + 1]
        assert segundo.text.strip() == "{{ ud_item }}", "se esperaba una pareja de {{ ud_item }}"
        _set_paragraph_text(primero, "{%p for x in list_uf_items %}")
        _set_paragraph_text(segundo, "{{ x }}")
        _nuevo_parrafo_tras(segundo, "{%p endfor %}")


def _borrar_lista_fija_ud_c1(doc):
    patron = re.compile(r"^UD \d+\.\s*\{\{ ud\d+_titulo_c1 \}\}\s*$")
    to_delete = [p for p in doc.paragraphs if patron.match(p.text.strip())]
    _borrar_parrafos(to_delete)


def _reemplazar_tabla_secuenciacion(doc):
    """Sustituye la tabla estática RA×UD (7 columnas de RA fijas) por el
    marcador [[TABLA_SECUENCIACION]] — generador_pd_suficiente_tpl.py la
    reconstruye con python-docx con las columnas/filas reales del módulo."""
    tabla = None
    for t in doc.tables:
        primera_fila_texto = " ".join(c.text for c in t.rows[0].cells)
        if "RESULTADOS DE APRENDIZAJE" in primera_fila_texto:
            tabla = t
            break
    if tabla is None:
        return

    ancla = doc.paragraphs[0]
    new_p = OxmlElement('w:p')
    tabla._tbl.addprevious(new_p)
    marcador = Paragraph(new_p, ancla._parent)
    marcador.add_run("[[TABLA_SECUENCIACION]]")
    tabla._tbl.getparent().remove(tabla._tbl)


def _dinamizar_seccion_c2(doc):
    """Sustituye el bloque estático "RA 1. {{ ra1_texto_c2 }}...RA 7...con
    sus CE hardcodeados" (justo antes de "C3. CRITERIOS DE CALIFICACIÓN")
    por un bucle dinámico sobre list_c2."""
    paras = doc.paragraphs
    idx_inicio = next((i for i, p in enumerate(paras) if p.text.strip() == "RA 1. {{ ra1_texto_c2 }}"), None)
    idx_fin = next((i for i, p in enumerate(paras) if p.text.strip() == "C3. CRITERIOS DE CALIFICACIÓN"), None)
    if idx_inicio is None or idx_fin is None or idx_fin <= idx_inicio:
        return

    ancla = paras[idx_inicio]
    _set_paragraph_text(ancla, "{%p for line in list_c2 %}")
    p2 = _nuevo_parrafo_tras(ancla, "{{ line }}")
    _nuevo_parrafo_tras(p2, "{%p endfor %}")

    # Borrar el resto del bloque estático original (de idx_inicio+1 a idx_fin-1)
    _borrar_parrafos(paras[idx_inicio + 1:idx_fin])


def _dinamizar_seccion_c3(doc):
    """Sustituye las líneas de porcentajes fijos de C3 (después de
    {{ texto_criterios_calificacion }}, antes de "En caso de no obtener...")
    por el marcador [[TABLA_INSTRUMENTOS_C3]]."""
    paras = doc.paragraphs
    idx_texto = next((i for i, p in enumerate(paras) if "{{ texto_criterios_calificacion }}" in p.text), None)
    idx_fin = next(
        (i for i, p in enumerate(paras) if p.text.strip().startswith("En caso de no obtener una evaluación positiva")),
        None,
    )
    if idx_texto is None or idx_fin is None or idx_fin <= idx_texto:
        return

    marcador = _nuevo_parrafo_tras(paras[idx_texto], "[[TABLA_INSTRUMENTOS_C3]]")
    # Borrar el resto del bloque estático (entre el nuevo marcador y "En caso de no obtener...")
    paras2 = doc.paragraphs
    idx_marcador = next(i for i, p in enumerate(paras2) if p.text.strip() == "[[TABLA_INSTRUMENTOS_C3]]")
    idx_fin2 = next(i for i, p in enumerate(paras2) if p.text.strip().startswith("En caso de no obtener una evaluación positiva"))
    _borrar_parrafos(paras2[idx_marcador + 1:idx_fin2])


def _dinamizar_seccion_h(doc):
    patron = re.compile(r"^RA \d+\.\s*\{\{ ra\d+_texto_c2 \}\}\s*$")
    paras = doc.paragraphs
    indices = [i for i, p in enumerate(paras) if patron.match(p.text.strip())]
    if not indices:
        return
    idx_inicio = indices[0]
    ancla = paras[idx_inicio]
    _set_paragraph_text(ancla, "{%p for ra in list_ras_h %}")
    p2 = _nuevo_parrafo_tras(ancla, "{{ ra }}")
    _nuevo_parrafo_tras(p2, "{%p endfor %}")
    _borrar_parrafos([paras[i] for i in indices[1:]])

    for p in doc.paragraphs:
        if "consta de seis Resultados de aprendizaje" in p.text:
            _set_paragraph_text(p, p.text.replace("seis", "{{ n_ra }}"))
            break


def _dinamizar_seccion_j(doc):
    paras = doc.paragraphs
    idx_inicio = next((i for i, p in enumerate(paras) if p.text.strip() == "Pizarra y proyector."), None)
    idx_fin = next(
        (i for i, p in enumerate(paras) if p.text.strip() == "Herramientas y equipos de medida del laboratorio."),
        None,
    )
    if idx_inicio is None or idx_fin is None:
        return

    ancla = paras[idx_inicio]
    _set_paragraph_text(ancla, "{%p for h in list_herramientas %}")
    p2 = _nuevo_parrafo_tras(ancla, "{{ h }}")
    _nuevo_parrafo_tras(p2, "{%p endfor %}")
    _borrar_parrafos(paras[idx_inicio + 1:idx_fin + 1])


def _dinamizar_seccion_n(doc):
    for p in doc.paragraphs:
        if p.text.strip().startswith("Debido a la dependencia que existe"):
            _set_paragraph_text(p, "{{ textos_pd_metodologia_labor_coordinada }}")
            return


def main():
    doc = Document(TEMPLATE_PATH)
    _quitar_anotaciones_sueltas(doc)
    _dinamizar_listas_ud_item(doc)
    _borrar_lista_fija_ud_c1(doc)
    _reemplazar_tabla_secuenciacion(doc)
    _dinamizar_seccion_c2(doc)
    _dinamizar_seccion_c3(doc)
    _dinamizar_seccion_h(doc)
    _dinamizar_seccion_j(doc)
    _dinamizar_seccion_n(doc)
    _quitar_comentarios_rf(doc)
    doc.save(TEMPLATE_PATH)
    print(f"[OK] Plantilla corregida y guardada en: {TEMPLATE_PATH}")


if __name__ == "__main__":
    main()
