import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

def generate_ud(ud_data, info_modulo, out_docx, out_pdf):
    doc = Document()
    
    title = doc.add_heading(f"Unidad Didáctica {ud_data.get('id_ud', '')}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading(ud_data.get("desc_ud", "Sin título"), level=1)
    
    doc.add_paragraph(f"Módulo: {info_modulo.get('modulo', '')}")
    doc.add_paragraph(f"Ciclo: {info_modulo.get('ciclo', '')}")
    doc.add_paragraph(f"Horas asignadas: {ud_data.get('horas_ud', 0)} h")
    
    doc.add_heading("Intención Educativa / Contextualización", level=2)
    doc.add_paragraph(ud_data.get('Intencion_Educativa', 'No especificada'))
    
    doc.add_heading("Temporización", level=2)
    doc.add_paragraph(ud_data.get('Temporizacion', 'No especificada'))
    
    doc.add_heading("Agrupamientos", level=2)
    doc.add_paragraph(ud_data.get('Agrupamientos', 'No especificados'))
    
    doc.add_heading("Transversalidad", level=2)
    doc.add_paragraph(ud_data.get('Transversalidad', 'No especificada'))

    doc.save(out_docx)
    try:
        convert(os.path.abspath(out_docx), os.path.abspath(out_pdf))
    except Exception as e:
        print(f"Error converting to PDF: {e}")
        if os.path.exists(out_pdf):
            os.remove(out_pdf)

def generate_tarea(tarea_data, info_modulo, out_docx, out_pdf):
    doc = Document()
    
    title = doc.add_heading(f"Tarea Competencial: {tarea_data.get('ID', tarea_data.get('id_act', ''))}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading(tarea_data.get("Nombre_Tarea", "Sin título"), level=1)
    
    doc.add_paragraph(f"Módulo: {info_modulo.get('modulo', '')}")
    
    doc.add_heading("Scenario / Briefing", level=2)
    doc.add_paragraph(tarea_data.get('Briefing', 'No especificado'))
    
    doc.add_heading("Pasos a seguir", level=2)
    doc.add_paragraph(tarea_data.get('Pasos', 'No especificados'))
    
    doc.add_heading("Evidencias de Entrega", level=2)
    doc.add_paragraph(tarea_data.get('Evidencias', 'No especificadas'))
    
    doc.add_heading("Formato y Plazo de Entrega", level=2)
    doc.add_paragraph(tarea_data.get('Entrega', 'No especificados'))

    doc.save(out_docx)
    try:
        convert(os.path.abspath(out_docx), os.path.abspath(out_pdf))
    except Exception as e:
        print(f"Error converting to PDF: {e}")
        if os.path.exists(out_pdf):
            os.remove(out_pdf)
