"""
Generador PD Mínima / Resumen (Vía A) — usa plantilla DOCX con docxtpl.

Plantilla: backend/templates/modelo_pd_fp-.docx
Destinatario: Alumnado (resumen de 1 página)

Uso:
    from generador_pd_minima_tpl import generate
    generate(data, out_docx, out_pdf)
"""

import os
from docxtpl import DocxTemplate
from helpers_catalogo import build_ra_desc_map, build_ud_desc_map, resolve_ra_desc, resolve_ud_desc

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), 'templates', 'modelo_pd_fp-.docx')


def _build_context(data: dict) -> dict:
    """
    Construye el contexto Jinja2 para la PD- (resumen alumnado).
    """
    modulo = data.get("modulo", "el módulo profesional")
    ciclo = data.get("ciclo", "el ciclo formativo")
    curso_ciclo = data.get("curso_ciclo", "")
    nivel = data.get("nivel", "")
    codificado = " ".join(p for p in [curso_ciclo, nivel] if p)
    curso = data.get("curso_academico", "")
    if codificado:
        curso = f"{curso} [{codificado}]" if curso else f"[{codificado}]"
    df_ra = data.get("df_ra", [])
    df_ud = data.get("df_ud", [])
    config = data.get("config_contexto", {})

    # Construir mapas de descripciones desde catálogo
    ra_desc_map = build_ra_desc_map(data)
    ud_desc_map = build_ud_desc_map(data)

    context = {
        "modulo": modulo,
        "ciclo": ciclo,
        "curso_academico": curso,
    }

    # ── H1 1: Perfil profesional ───────────────────────────────────────
    texto_perfil = config.get("minima_perfil_profesional", "")
    if not texto_perfil:
        texto_perfil = (
            f"La formación de este módulo ({modulo}) contribuye a tu perfil profesional "
            f"proporcionándote las competencias necesarias para desarrollar tu actividad "
            f"en el ámbito del {ciclo}."
        )
    context["texto_perfil_profesional"] = texto_perfil

    # ── H1 2: Competencia general ──────────────────────────────────────
    texto_comp = config.get("minima_competencia_general", "")
    if not texto_comp:
        texto_comp = (
            f"La competencia general de este título te permitirá desempeñar las funciones "
            f"propias del perfil profesional asociado al {ciclo}, aplicando los conocimientos "
            f"y habilidades adquiridos en {modulo}."
        )
    context["texto_competencia_general"] = texto_comp

    # --- H1 3: Resultados de Aprendizaje (lista), con % y las UD que lo cubren ──
    list_ras = []
    for ra in df_ra:
        id_str = str(ra.get('id_ra', '')).strip().rstrip('.')
        prefix = "" if id_str.upper().startswith("RA") else "RA"
        ra_id_full = f"{prefix}{id_str}"
        desc = resolve_ra_desc(ra, ra_desc_map)
        try:
            peso = int(float(ra.get('peso_ra', 0) or 0))
        except (ValueError, TypeError):
            peso = 0

        # UDs relacionadas con este RA (misma info que la Matriz RA↔UD)
        uds_rel = []
        for ud in df_ud:
            try:
                val = float(ud.get(ra_id_full, 0) or 0)
            except (ValueError, TypeError):
                val = 0
            if val > 0:
                ud_id = str(ud.get('id_ud', '')).strip()
                horas = int(float(ud.get('horas_ud', 0) or 0))
                uds_rel.append(f"{ud_id} ({horas}h) - {int(val)}%")

        # Línea del RA y, como elemento de lista aparte (para poder darle
        # su propio formato de párrafo: indentado y en cursiva), la línea
        # con las UD relacionadas — siempre empieza por "UDxx (" para poder
        # reconocerla después del render, ver _sangrar_lineas_relacion_ud().
        list_ras.append(f"{ra_id_full}. ({peso}%) {desc}")
        if uds_rel:
            list_ras.append(", ".join(uds_rel))
    context["list_ras"] = list_ras

    # --- H1 4: Contenidos / UDs (lista), con horas ──────────────────────
    list_uds = []
    for ud in df_ud:
        id_str = str(ud.get('id_ud', '')).strip().rstrip('.')
        prefix = "" if id_str.upper().startswith("UD") else "UD"
        desc = resolve_ud_desc(ud, ud_desc_map)
        horas = int(float(ud.get('horas_ud', 0) or 0))
        list_uds.append(f"{prefix}{id_str}. ({horas}h) {desc}")
    context["list_uds"] = list_uds

    # --- % Ponderación por trimestres (para la nota final) ──────────────
    info_mod_pond = data.get("info_modulo") or {}
    context["pond_1t"] = info_mod_pond.get("pond_1t", 30)
    context["pond_2t"] = info_mod_pond.get("pond_2t", 30)
    context["pond_3t"] = info_mod_pond.get("pond_3t", 40)

    # --- H1 5: Criterios de calificación (% instrumentos por trimestre) ---
    instrumentos_pct = data.get("instrumentos_pct_trimestre") or []
    if not instrumentos_pct:
        # Misma semilla por defecto que el bloque "% Instrumentos de evaluación"
        # de Programación › Contexto › Identificación, para que PD- y la app
        # muestren siempre lo mismo mientras el profesor no lo edite.
        instrumentos_pct = [
            {"nombre": "Exámenes teóricos", "pct_1t": 30, "pct_2t": 20, "pct_3t": 10},
            {"nombre": "Exámenes prácticos", "pct_1t": 20, "pct_2t": 20, "pct_3t": 10},
            {"nombre": "Exposición y defensa proyecto", "pct_1t": 10, "pct_2t": 20, "pct_3t": 30},
            {"nombre": "Informes de ejercicios", "pct_1t": 20, "pct_2t": 30, "pct_3t": 40},
            {"nombre": "Cuaderno de tareas", "pct_1t": 20, "pct_2t": 10, "pct_3t": 10},
        ]
    list_instrumentos = [
        {
            "nombre": row.get("nombre", ""),
            "pct_1t": f"{row.get('pct_1t', 0)}%",
            "pct_2t": f"{row.get('pct_2t', 0)}%",
            "pct_3t": f"{row.get('pct_3t', 0)}%",
        }
        for row in instrumentos_pct
    ]
    context["list_instrumentos"] = list_instrumentos

    # ── H1 6: Recordad ─────────────────────────────────────────────────
    texto_recordatorio = config.get("minima_recordatorio", "")
    if not texto_recordatorio:
        texto_recordatorio = "Más del 15% de faltas de asistencia a clase implica la Pérdida del derecho a evaluación continua."
    context["texto_recordatorio"] = texto_recordatorio

    texto_final = config.get("minima_texto_final", "")
    if not texto_final:
        texto_final = "Tu situación es similar a la del resto del alumnado que ha obtenido su titulación. ¡ÁNIMO!"
    context["texto_final"] = texto_final

    return context


def _quitar_vinetas_listas(doc):
    """Las listas de RA y UD usan en la plantilla el estilo 'List Bullet'
    (numPr → numId=1 en styles.xml), que es lo que pinta el puntito. Al
    pasarlas a 'Normal' se pierde la numeración (Normal no tiene numPr) y
    se fija a mano la sangría uniforme pedida — sin viñeta en ningún ítem,
    ni en RA, ni en UD, ni en la línea de relación RA↔UD."""
    from docx.shared import Cm
    for p in doc.paragraphs:
        if p.style is not None and p.style.name == 'List Bullet':
            p.style = doc.styles['Normal']
            p.paragraph_format.left_indent = Cm(2.0)


def _cursivar_lineas_relacion_ud(doc):
    """Las líneas 'UDxx (Yh) - Z%, ...' que list_ras intercala tras cada RA
    se reconocen por empezar con 'UD' + dígito, y se marcan en cursiva —
    un párrafo aparte no puede llevar formato distinto al de la línea
    anterior si van dentro del mismo párrafo, así que esto solo funciona
    porque _build_context ya las separa en su propio ítem de lista."""
    import re
    patron = re.compile(r"^UD\d+\s*\(")
    for p in doc.paragraphs:
        if patron.match(p.text.strip()):
            for r in p.runs:
                r.italic = True


def _forzar_arial(doc):
    """Une el documento bajo una única tipografía (Arial) para que no queden
    mezclas con la fuente por defecto de la plantilla en los párrafos que no
    tocamos a mano (título, perfil profesional, listas de RA/UD, etc.)."""
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.name = 'Arial'
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = 'Arial'


def generate(data: dict, out_docx: str, out_pdf: str = None):
    """
    Genera la PD- (resumen alumnado) usando la plantilla DOCX.
    """
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(
            f"No se encontró la plantilla en: {TEMPLATE_PATH}. "
            f"Ejecuta 'python scripts/preparar_plantilla_pd_minima.py' para generarla."
        )

    tpl = DocxTemplate(TEMPLATE_PATH)
    context = _build_context(data)
    tpl.render(context)

    doc = tpl.docx
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # ── Márgenes de página: 2 izq., 1 arriba/abajo/derecha ───────────
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(1)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)

    # ── Insertar Tabla de Instrumentos ──────────────────────────────
    for p in doc.paragraphs:
        if "[[TABLA_INSTRUMENTOS]]" in p.text:
            p.text = p.text.replace("[[TABLA_INSTRUMENTOS]]", "")

            # Crear la tabla dinámicamente
            list_inst = context.get("list_instrumentos", [])
            table = doc.add_table(rows=1 + len(list_inst), cols=4)
            table.style = 'Table Grid'

            # 18 cm de ancho total = 21 cm (A4) - 2 cm izq. - 1 cm der.
            col_widths = [Cm(9.0), Cm(3.0), Cm(3.0), Cm(3.0)]
            for i, w in enumerate(col_widths):
                table.columns[i].width = w
            # table.columns[i].width por sí solo no basta: Word respeta el
            # ancho de cada celda (tcW) por encima del gridCol de la tabla,
            # así que hay que fijarlo también celda a celda, en todas las filas.
            for row in table.rows:
                for i, w in enumerate(col_widths):
                    row.cells[i].width = w

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Instrumento"
            hdr_cells[1].text = f"1er Tri. ({context['pond_1t']}%)"
            hdr_cells[2].text = f"2º Tri. ({context['pond_2t']}%)"
            hdr_cells[3].text = f"3er Tri. ({context['pond_3t']}%)"

            for i, cell in enumerate(hdr_cells):
                for hp in cell.paragraphs:
                    if i > 0:
                        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in hp.runs:
                        r.bold = True
                        r.font.name = 'Arial'
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(0, 0, 0)

            for i, instr in enumerate(list_inst):
                row_cells = table.rows[i + 1].cells
                row_cells[0].text = instr.get("nombre", "")
                row_cells[1].text = instr.get("pct_1t", "")
                row_cells[2].text = instr.get("pct_2t", "")
                row_cells[3].text = instr.get("pct_3t", "")

                for j, cell in enumerate(row_cells):
                    for cp in cell.paragraphs:
                        if j > 0:
                            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in cp.runs:
                            r.font.name = 'Arial'
                            r.font.size = Pt(9)
                            r.font.color.rgb = RGBColor(0, 0, 0)

            # Mover la tabla para que quede después de este párrafo
            p._p.addnext(table._tbl)

    _quitar_vinetas_listas(doc)
    _cursivar_lineas_relacion_ud(doc)
    _forzar_arial(doc)

    # ── Guardar DOCX (y PDF) ─────────────────────────────────────────
    tpl.save(out_docx)
