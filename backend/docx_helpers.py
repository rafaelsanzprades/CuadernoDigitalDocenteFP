# -*- coding: utf-8 -*-
"""
docx_helpers.py
Utilidades comunes para construir las versiones .docx "editables" de los
documentos que hoy solo existen en PDF (ReportLab). No son conversiones
docx->pdf (docx2pdf está descartado, falla en Linux/Cloud Run): son
documentos .docx nativos, construidos directamente con python-docx,
reutilizando los mismos datos que ya prepara routers/pdf.py.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def new_document(landscape: bool = False) -> Document:
    doc = Document()
    section = doc.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    for attr in ("left_margin", "right_margin"):
        setattr(section, attr, Cm(1.5))
    for attr in ("top_margin", "bottom_margin"):
        setattr(section, attr, Cm(1.5))
    return doc


def add_title(doc: Document, titulo: str, subtitulo: str = "") -> None:
    h = doc.add_heading(titulo, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitulo:
        p = doc.add_paragraph(subtitulo)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)


def add_meta_line(doc: Document, texto: str) -> None:
    p = doc.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)


def add_section_heading(doc: Document, texto: str) -> None:
    h = doc.add_heading(texto, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def _shade_cell(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table(doc: Document, headers: list, rows: list, col_widths_cm: list = None,
              header_bg: str = "E0E0E0", total_row_bg: str = None) -> "Document.tables":
    """rows: list[list[str]]. Si total_row_bg se indica, la última fila de `rows`
    se sombrea (para filas de totales)."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
        _shade_cell(hdr_cells[i], header_bg)

    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        is_total = total_row_bg and r_idx == len(rows) - 1
        for c_idx, val in enumerate(row):
            cells[c_idx].text = "" if val is None else str(val)
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if is_total:
                        run.font.bold = True
            if is_total:
                _shade_cell(cells[c_idx], total_row_bg)

    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def doc_to_bytes(doc: Document) -> bytes:
    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
