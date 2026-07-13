import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
try:
    from docx2pdf import convert
except ImportError:
    convert = None

def _add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h

def _add_toc(doc):
    doc.add_heading('Índice', level=1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    doc.add_page_break()

def generate(data, out_docx, out_pdf):
    doc = Document()
    
    # --- Portada ---
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Programación didáctica Aragón (BOA nº: 181 de 18 de septiembre de 2025)")
    run.font.size = Pt(24)
    run.font.bold = True
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run(data.get("modulo", "Módulo Profesional"))
    run2.font.size = Pt(18)
    
    doc.add_paragraph()
    
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.add_run(f"Ciclo Formativo: {data.get('ciclo', '')}\n")
    info_p.add_run(f"Departamento: {data.get('departamento', '')}\n")
    info_p.add_run(f"Curso Académico: {data.get('curso_academico', '')}\n")
    
    doc.add_page_break()
    
    _add_toc(doc)
    
    # A. INTRODUCCIÓN
    _add_heading(doc, "A. INTRODUCCIÓN", 1)
    doc.add_paragraph("La presente programación didáctica se ha elaborado conforme a la normativa vigente que establece el currículo del ciclo formativo en la Comunidad Autónoma, adaptándose al contexto socioeducativo del centro y a las características del alumnado.")

    # B. RESULTADOS DE APRENDIZAJE SUSCEPTIBLES DE SER ADQUIRIDOS EN LA FORMACIÓN EN EMPRESA
    _add_heading(doc, "B. RESULTADOS DE APRENDIZAJE SUSCEPTIBLES DE SER ADQUIRIDOS EN LA FORMACIÓN EN EMPRESA", 1)
    if data.get("is_dual", False):
        doc.add_paragraph("Este módulo PARTICIPA en la Formación en Empresa u Organismo Equiparado (FEOE) / Modalidad Dual.")
        doc.add_paragraph("Los resultados de aprendizaje que se podrán adquirir o completar en la empresa se determinarán en el plan de formación individualizado de cada alumno/a.")
    else:
        doc.add_paragraph("Este módulo NO se desarrolla en modalidad Dual (no tiene horas de FEOE asociadas).")

    # C. SECUENCIACIÓN Y ORGANIZACIÓN TEMPORAL DE LAS UNIDADES DIDÁCTICAS
    _add_heading(doc, "C. SECUENCIACIÓN Y ORGANIZACIÓN TEMPORAL DE LAS UNIDADES DIDÁCTICAS", 1)
    df_ud = data.get("df_ud", [])
    if df_ud:
        for ud in df_ud:
            doc.add_paragraph(f"UD {ud.get('id_ud', '')}: {ud.get('desc_ud', '')} ({ud.get('horas_ud', 0)} horas)", style='List Number')
    else:
        doc.add_paragraph("No hay unidades didácticas definidas.")

    # C1. CONTENIDOS
    _add_heading(doc, "C1. CONTENIDOS", 2)
    doc.add_paragraph("Los contenidos están integrados y secuenciados dentro de las Unidades Didácticas presentadas en el apartado anterior.")

    # C2. CRITERIOS DE EVALUACIÓN
    _add_heading(doc, "C2. CRITERIOS DE EVALUACIÓN", 2)
    doc.add_paragraph("Se aplicarán los criterios de evaluación normativos definidos en el RD y la Orden del Título, asociados a cada resultado de aprendizaje.")

    # C3. CRITERIOS DE CALIFICACIÓN
    _add_heading(doc, "C3. CRITERIOS DE CALIFICACIÓN", 2)
    redondeo = data.get("config_redondeo", {})
    doc.add_paragraph(f"Nota mínima para aprobar: {redondeo.get('nota_aprobado', 5.0)}")
    doc.add_paragraph(f"Redondeo al alza desde: {redondeo.get('umbral_redondeo', 4.5)}")

    # D. PRINCIPIOS METODOLÓGICOS
    _add_heading(doc, "D. PRINCIPIOS METODOLÓGICOS", 1)
    metodologias = data.get("metodologias_seleccionadas", [])
    if metodologias:
        doc.add_paragraph(f"Metodologías aplicadas: {', '.join(metodologias)}.")
    doc.add_paragraph(data.get("texto_metodologia_libre", "Las actividades tendrán un carácter fundamentalmente práctico."))

    # E. EVALUACIÓN INICIAL
    _add_heading(doc, "E. EVALUACIÓN INICIAL", 1)
    doc.add_paragraph("Durante las primeras semanas del curso se realizará una evaluación inicial para conocer el nivel de partida del alumnado y adaptar el proceso de enseñanza-aprendizaje.")

    # E1. ATENCIÓN A LAS DIFERENCIAS INDIVIDUALES
    _add_heading(doc, "E1. ATENCIÓN A LAS DIFERENCIAS INDIVIDUALES", 2)
    inclusion = data.get("medidas_inclusion", [])
    if inclusion:
        doc.add_paragraph(f"Medidas adoptadas: {', '.join(inclusion)}.")
    doc.add_paragraph(data.get("texto_inclusion_libre", ""))

    # F. PROCEDIMIENTOS E INSTRUMENTOS DE EVALUACIÓN
    _add_heading(doc, "F. PROCEDIMIENTOS E INSTRUMENTOS DE EVALUACIÓN", 1)
    instrumentos = data.get("instrumentos_seleccionados", [])
    if instrumentos:
        doc.add_paragraph(f"Instrumentos principales: {', '.join(instrumentos)}.")
    else:
        doc.add_paragraph("Instrumentos variados según el tipo de actividad (pruebas teóricas, rúbricas, trabajos prácticos).")

    # G. ACTIVIDADES DE RECUPERACIÓN Y REFUERZO
    _add_heading(doc, "G. ACTIVIDADES DE RECUPERACIÓN Y REFUERZO", 1)
    doc.add_paragraph("Se diseñarán actividades de refuerzo para aquel alumnado que no haya alcanzado los resultados de aprendizaje o necesite consolidar conocimientos.")

    # G1. PLAN DE RECUPERACIÓN
    _add_heading(doc, "G1. PLAN DE RECUPERACIÓN", 2)
    doc.add_paragraph("El alumnado que no supere alguna evaluación dispondrá de mecanismos de recuperación (pruebas objetivas o entrega de trabajos prácticos) antes de la evaluación final.")

    # H. RESULTADOS DE APRENDIZAJE
    _add_heading(doc, "H. RESULTADOS DE APRENDIZAJE", 1)
    df_ra = data.get("df_ra", [])
    if df_ra:
        for ra in df_ra:
            doc.add_paragraph(f"RA {ra.get('id_ra', '')}: {ra.get('desc_ra', '')}", style='List Bullet')
    else:
        doc.add_paragraph("No hay resultados de aprendizaje definidos.")

    # I. PLAN DE APLICACIÓN DE LOS DESDOBLES, EN SU CASO
    _add_heading(doc, "I. PLAN DE APLICACIÓN DE LOS DESDOBLES, EN SU CASO", 1)
    doc.add_paragraph("En caso de disponer de horas de desdoble, se dedicarán a prácticas de taller/laboratorio para garantizar la seguridad y una atención más individualizada.")

    # J. MATERIALES Y RECURSOS DIDÁCTICOS
    _add_heading(doc, "J. MATERIALES Y RECURSOS DIDÁCTICOS", 1)
    recursos = data.get("recursos_espacios", [])
    if recursos:
        doc.add_paragraph(f"Recursos y espacios: {', '.join(recursos)}.")
    else:
        doc.add_paragraph("Equipamiento del aula técnica, recursos digitales y bibliografía recomendada.")

    # K. ACTIVIDADES COMPLEMENTARIAS Y EXTRAESCOLARES
    _add_heading(doc, "K. ACTIVIDADES COMPLEMENTARIAS Y EXTRAESCOLARES", 1)
    extraescolares = data.get("actividades_complementarias", [])
    if extraescolares:
        doc.add_paragraph(f"Actividades propuestas: {', '.join(extraescolares)}.")
    else:
        doc.add_paragraph("A determinar por el departamento a lo largo del curso.")

    # L. MEDIDAS COMPLEMENTARIAS EN PROYECTOS O BILINGÜES, EN SU CASO
    _add_heading(doc, "L. MEDIDAS COMPLEMENTARIAS EN PROYECTOS O BILINGÜES, EN SU CASO", 1)
    transversales = data.get("elementos_transversales", [])
    if transversales:
        doc.add_paragraph(f"Elementos transversales y proyectos: {', '.join(transversales)}.")
    else:
        doc.add_paragraph("No se contemplan medidas específicas adicionales en este apartado.")

    # M. MECANISMOS DE SEGUIMIENTO Y VALORACIÓN
    _add_heading(doc, "M. MECANISMOS DE SEGUIMIENTO Y VALORACIÓN", 1)
    doc.add_paragraph("Al finalizar cada trimestre y a final de curso, el equipo docente analizará el grado de cumplimiento de esta programación, aplicando medidas correctoras si fuera necesario.")

    # N. PLAN DE CONTINGENCIA
    _add_heading(doc, "N. PLAN DE CONTINGENCIA", 1)
    contingencia = data.get("medidas_contingencia", [])
    if contingencia:
        doc.add_paragraph(f"Medidas: {', '.join(contingencia)}.")
    doc.add_paragraph(data.get("texto_contingencia_libre", ""))

    # ANEXO NORMATIVO
    doc.add_page_break()
    _add_heading(doc, "ANEXO: CUMPLIMIENTO NORMATIVO Y ESTRUCTURA", 1)
    
    p_norma = doc.add_paragraph()
    p_norma.add_run("Actualizado al BOA nº: 181 de 18 de septiembre de 2025, página 23, título Treinta, “Los apartados 1 y 2 del artículo 100 (del BOE nº: 109 de 6 de junio de 2024 DECRETO 91/2024, de 5 de junio, del Gobierno de Aragón por el que se establece la Ordenación de la Formación Profesional del Grado D y del Grado E en la Comunidad Autónoma de Aragón) quedan redactados como sigue”, su apartado “2. La programación didáctica de cada módulo y, en su caso, ámbito y Proyecto, estará compuesta, al menos, por los siguientes elementos”").italic = True
    
    doc.add_paragraph("Para facilitar la lectura, se han simplificado los epígrafes en esta programación; no obstante, la misma da respuesta íntegra a los apartados exigidos por la norma, cuya redacción completa se detalla a continuación:")
    
    norma_items = [
        "a) Introducción en la que se incluya, entre otras cosas, el marco normativo y su contextualización de acuerdo con el entorno del centro docente.",
        "b) Los resultados de aprendizaje susceptibles de ser adquiridos en la formación en empresa u organismo equiparado, justificando los motivos por los cuales se dualizan.",
        "c) Secuenciación y organización temporal de las unidades didácticas que se vayan a impartir en el centro docente asociadas a los contenidos, criterios de evaluación, resultados de aprendizaje, objetivos generales y competencias profesionales y para la empleabilidad.",
        "d) Los principios metodológicos a desarrollar en el módulo o, en su caso, ámbito y Proyecto. Si el módulo o, en su caso, ámbito, se trabaja por proyectos y/o retos, se deberá indicar la metodología utilizada, vinculándolo con el entorno productivo.",
        "e) Las características de la evaluación inicial, criterios para su valoración, así como consecuencias de sus resultados en la programación didáctica y el diseño de los instrumentos de evaluación.",
        "f) Los procedimientos e instrumentos de evaluación, con la participación del/ de la tutor/a de empresa u organismo equiparado en los módulos dualizados. Asimismo, se deberá incluir su vinculación con los criterios de evaluación y los criterios de calificación del módulo o, en su caso, ámbito y Proyecto, incluyendo los utilizados para el alumnado que pierde el derecho a la evaluación continua, teniendo en cuenta que se deben utilizar instrumentos de evaluación variados.",
        "g) Las actividades de recuperación y refuerzo previstas para el alumnado que tenga que presentarse a la segunda convocatoria de evaluación final, la atención a las diferencias individuales y el plan de recuperación del módulo o, en su caso, ámbito pendiente. Estas actividades de recuperación y refuerzo deben fundamentarse en los resultados de aprendizaje en los que el alumnado haya tenido dificultades. Asimismo, se incluirán las actividades de ampliación para el alumnado que las precise.",
        "h) Calificación mínima que se debe alcanzar en cada resultado de aprendizaje para que puedan compensarse entre ellos, de acuerdo con lo establecido en el Proyecto curricular.",
        "i) El plan de aplicación de los desdobles, en su caso, de acuerdo con lo establecido en el Proyecto curricular.",
        "j) Los materiales y recursos didácticos que se vayan a utilizar, teniendo en cuenta que deben ser diversos y coherentes con la metodología especificada y con la planificación del uso de espacios y equipamientos, entre otros.",
        "k) Las actividades complementarias y extraescolares programadas, concretando la incidencia y los objetivos de las mismas en la evaluación del alumnado en relación con los resultados de aprendizaje.",
        "l) En su caso, las medidas complementarias que se plantean para el tratamiento de los módulos dentro de proyectos o itinerarios bilingües.",
        "m) Los mecanismos de seguimiento y valoración de la impartición del módulo o, en su caso, ámbito y Proyecto, que permita potenciar los resultados positivos y subsanar las deficiencias que se hayan detectado, teniendo en cuenta lo establecido en el Plan de mejora.",
        "n) Un plan de contingencia con las actividades que realizarán las personas en formación ante circunstancias excepcionales que afecten al desarrollo normal de la actividad docente en el módulo o, en su caso, ámbito y Proyecto, durante un período prologando de tiempo."
    ]
    
    for item in norma_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Pt(20)

    # Force TOC update on open
    try:
        element = doc.settings.element
        updateFields = OxmlElement('w:updateFields')
        updateFields.set(qn('w:val'), 'true')
        element.append(updateFields)
    except Exception:
        pass

    # Save DOCX
    doc.save(out_docx)
    
    # Convert to PDF
    try:
        abs_docx = os.path.abspath(out_docx)
        abs_pdf = os.path.abspath(out_pdf)
        if convert:
            convert(abs_docx, abs_pdf)
        else:
            print("docx2pdf not available, skipping PDF conversion.")
    except Exception as e:
        print(f"Error converting to PDF: {e}")
        if os.path.exists(out_pdf):
            os.remove(out_pdf)
