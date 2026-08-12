"""
Helpers de tablas python-docx compartidos entre generadores de PD (PD- y PD=).

Construir la tabla directamente con python-docx (en vez de una tabla estática
en la plantilla docxtpl) es lo que permite que tenga tantas filas como haga
falta y que el ancho de columna se controle de forma fiable — ver el propio
comentario de insertar_tabla_instrumentos sobre el bug de table.columns[i].width.
"""

from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TEXTO_ESTIMACION_PLANIFICACION = (
    "Esta previsión estima que las horas lectivas de cada semana se cumplen de forma "
    "proporcional a los días festivos o no lectivos que contenga, a falta de que se "
    "concreten el calendario oficial del curso académico y la distribución definitiva "
    "del horario semanal por días. Es una aproximación orientativa, no una previsión exacta."
)


def aplicar_bordes_rejilla(table):
    """Aplica bordes de cuadrícula a `table` a mano (equivalente al estilo
    'Table Grid') en vez de asignar table.style — no todas las plantillas
    tienen ese estilo activado (la del BOA no lo trae), así que confiar en
    el nombre de estilo rompe con KeyError en cuanto cambia la plantilla."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tbl_pr.append(borders)


def insertar_tabla_instrumentos(doc, paragraph, list_instrumentos, pond_1t, pond_2t, pond_3t):
    """
    Inserta, justo después de `paragraph`, la tabla "Instrumento | 1er/2º/3er
    Tri." (9+3+3+3cm, columnas 2-4 centradas) y elimina `paragraph` (se asume
    que solo contenía el marcador de posición, ya vacío en este punto) para
    que no quede una línea en blanco antes de la tabla.
    """
    table = doc.add_table(rows=1 + len(list_instrumentos), cols=4)
    aplicar_bordes_rejilla(table)

    # 18 cm de ancho total = 21 cm (A4) - 2 cm izq. - 1 cm der.
    col_widths = [Cm(9.0), Cm(3.0), Cm(3.0), Cm(3.0)]
    for i, w in enumerate(col_widths):
        table.columns[i].width = w
    # table.columns[i].width por sí solo no basta: Word respeta el ancho de
    # cada celda (tcW) por encima del gridCol de la tabla, así que hay que
    # fijarlo también celda a celda, en todas las filas.
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Instrumento"
    hdr_cells[1].text = f"1er Tri. ({pond_1t}%)"
    hdr_cells[2].text = f"2º Tri. ({pond_2t}%)"
    hdr_cells[3].text = f"3er Tri. ({pond_3t}%)"

    for i, cell in enumerate(hdr_cells):
        for hp in cell.paragraphs:
            if i > 0:
                hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in hp.runs:
                r.bold = True
                r.font.name = 'Arial'
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0, 0, 0)

    for i, instr in enumerate(list_instrumentos):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = instr.get("nombre", "")
        row_cells[1].text = instr.get("pct_1t", "")
        row_cells[2].text = instr.get("pct_2t", "")
        row_cells[3].text = instr.get("pct_3t", "")

        for j, cell in enumerate(row_cells):
            for cp in cell.paragraphs:
                if j > 0:
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in cp.runs:
                    r.font.name = 'Arial'
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0, 0, 0)

    paragraph._p.addnext(table._tbl)
    paragraph._p.getparent().remove(paragraph._p)
    return table


def _shade_cell(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def insertar_tabla_planificacion(doc, df_ud, info_fechas, horario, calendar_notes):
    """
    Añade al final de `doc`, en una nueva sección apaisada, la previsión de
    planificación mensual: Tri. | UD | Prev. total | Prev. por mes, con la
    fila FEOE colocada por orden cronológico (no siempre al final).

    A diferencia de /agenda?tab=planificacion (PlanificacionMensualTab.tsx),
    que muestra Prv/Imp/%Imp/Prv porque es una herramienta de seguimiento
    día a día, esta página del PD es solo una previsión — un documento que
    normalmente se entrega antes de empezar a impartir el módulo, así que no
    tiene sentido mostrar una columna "impartido" (a petición de Rafael:
    "es una previsión... solo se debe mostrar una columna la llamada Prev.").
    Por eso una sola columna por mes, en vez de la pareja Prv/Imp.

    Recalcula el reparto con planning_generator.generate_planning() — el
    port a Python de planningGenerator.ts — en vez de leer el df_sgmt
    guardado en curso_data: esa página nunca lee el snapshot guardado, lo
    recalcula en cada render a partir de df_ud + info_fechas + horario +
    calendar_notes, así que el snapshot puede estar desactualizado respecto
    a esos datos de origen. Recalcular aquí es lo único que garantiza que el
    PD coincide con lo que ve el profesorado en la app.

    La sección apaisada es necesaria porque la tabla tiene 13 columnas
    (3 fijas + 10 meses); doc.sections[0] no sirve porque solo controla
    la orientación de todo el documento, no de un tramo — de ahí
    doc.add_section(WD_SECTION.NEW_PAGE) + orientación propia de esa sección.
    """
    from planning_generator import generate_planning, MESES_DISPLAY as meses_display

    df_ud = df_ud or []
    filas = generate_planning(df_ud, info_fechas, horario, calendar_notes)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    doc.add_heading("Previsión de planificación mensual", level=1)

    p_nota = doc.add_paragraph(TEXTO_ESTIMACION_PLANIFICACION)
    for r in p_nota.runs:
        r.italic = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    N_FIXED = 3  # Tri., UD, Prev.
    n_cols = N_FIXED + len(meses_display)

    ud_desc_map = {str(ud.get("id_ud", "")): ud.get("desc_ud", "") for ud in df_ud}

    def _ud_label(fila):
        uid = str(fila.get("id_ud", ""))
        if uid == "Sin docencia":
            return uid
        desc = fila.get("desc_ud") or ud_desc_map.get(uid, "")
        return f"{uid}. {desc}" if desc else uid

    table = doc.add_table(rows=1 + len(filas), cols=n_cols)
    aplicar_bordes_rejilla(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    HEADER_BG = "F0F0F0"

    def set_cell(cell, text, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.CENTER):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text))
        run.bold = bold
        run.font.name = 'Arial'
        run.font.size = Pt(size)

    headers = ["Tri.", "UD. Unidad didáctica", "Prev."] + list(meses_display)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, size=8, align=(WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER))
        _shade_cell(cell, HEADER_BG)

    for r_idx, fila in enumerate(filas):
        r = 1 + r_idx
        ev = fila.get("ev")
        horas_ud = fila.get("horas_ud", 0) or 0

        set_cell(table.rows[r].cells[0], f"{ev}ª" if ev else "", bold=True, size=8)
        set_cell(table.rows[r].cells[1], _ud_label(fila), bold=True, size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(table.rows[r].cells[2], horas_ud if horas_ud else "", bold=True, size=8)
        for mi, m in enumerate(meses_display):
            prv = int(fila.get(f"{m}_Prv", 0) or 0)
            set_cell(table.rows[r].cells[N_FIXED + mi], prv if prv else "", size=7)

    col_widths = [Cm(1.2), Cm(8.5), Cm(1.5)] + [Cm(1.5)] * len(meses_display)
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w
