"""
Generador PD Suficiente / BOA (Vía A) — usa plantilla DOCX con docxtpl.

Plantilla: backend/templates/modelo_pd_fp=.docx
Estructura: 14 secciones A-N (modelo BOA/Aragón)

Uso:
    from generador_pd_suficiente_tpl import generate
    generate(data, out_docx, out_pdf)
"""

import os
from docxtpl import DocxTemplate
from helpers_catalogo import (
    build_ra_desc_map, build_ud_desc_map, build_ce_desc_map,
    resolve_ra_desc, resolve_ud_desc, resolve_ce_desc, resolve_recursos,
)
from helpers_pd_tablas import insertar_tabla_instrumentos

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), 'templates', 'modelo_pd_fp=.docx')

DEFAULT_HERRAMIENTAS = [
    "Pizarra y proyector.",
    "Ordenadores e impresora para realización de las prácticas y/o sus informes.",
    "Software ofimático y aplicaciones del módulo.",
    "Herramientas y equipos de medida del laboratorio.",
]

DEFAULT_TEXTO_COORDINACION = (
    "Debido a la dependencia que existe entre los módulos y su profesorado, se ha optado "
    "por establecer reuniones periódicas de coordinación del equipo docente."
)


def _ra_id_full(ra: dict) -> str:
    id_str = str(ra.get('id_ra', '')).strip().rstrip('.')
    prefix = "" if id_str.upper().startswith("RA") else "RA"
    return f"{prefix}{id_str}"


def _build_context(data: dict) -> dict:
    """
    Construye el contexto Jinja2 para la PD= (Suficiente/BOA).

    data: {
        "modulo": "Maquinas Electricas",
        "ciclo": "Tecnico IEA",
        "departamento": "Electricidad",
        "curso_academico": "2025/2026",
        "horas_totales": 167,
        "horas_semanales": 5,
        "regimen": "diurno",
        "curso_ciclo": "primero",
        "df_ra": [...],
        "df_ud": [...],
        "config_contexto": {...}
    }
    """
    info_mod = data.get("info_modulo") or {}
    modulo = data.get("modulo", "el módulo profesional")
    codigo = info_mod.get("codigo", "")
    if codigo:
        modulo = f"{codigo} - {modulo}"
    ciclo = data.get("ciclo", "el ciclo formativo")
    departamento = data.get("departamento", "")
    nivel = data.get("nivel", "")
    codificado = " ".join(p for p in [data.get("curso_ciclo", ""), nivel] if p)
    curso_academico_raw = data.get("curso_academico", "")
    curso = f"{codificado} - {curso_academico_raw}" if (codificado and curso_academico_raw) else (codificado or curso_academico_raw)
    horas_totales = data.get("horas_totales", "")
    horas_semanales = data.get("horas_semanales", "")
    regimen = data.get("regimen", "diurno")
    curso_ciclo = data.get("curso_ciclo", "primero")
    df_ra = data.get("df_ra", [])
    df_ud = data.get("df_ud", [])
    df_ce = data.get("df_ce", [])
    df_sgmt = data.get("df_sgmt", [])
    config = data.get("config_contexto", {})

    # Mapa id_ud -> evaluación (1/2/3) calculado en Planificación mensual
    # (frontend/src/utils/planningGenerator.ts) a partir de la fecha en la
    # que termina cada UD; no hay ningún otro campo en la app que lo derive.
    ev_map = {str(row.get("id_ud", "")): row.get("ev") for row in df_sgmt}

    # Construir mapas de descripciones desde catálogo (df_ra/df_ud/df_ce ->
    # curriculo_data -> catálogo oficial BOA/BOE resuelto en routers/pdf.py)
    ra_desc_map = build_ra_desc_map(data)
    ud_desc_map = build_ud_desc_map(data)
    ce_desc_map = build_ce_desc_map(data)

    context = {
        "modulo": modulo,
        "ciclo": ciclo,
        "departamento": departamento,
        "curso_academico": curso,
        "centro": info_mod.get("centro", ""),
    }

    # ── SECCIÓN A: INTRODUCCIÓN ────────────────────────────────────────
    texto_intro = config.get("texto_introduccion", "")
    if not texto_intro:
        texto_intro = (
            f"Programación didáctica del módulo profesional {modulo}, "
            f"perteneciente al {curso_ciclo} curso del ciclo formativo de grado "
            f"medio {ciclo}, que se imparte en régimen {regimen}, con una duración "
            f"de {horas_totales} periodos lectivos a lo largo del curso académico, "
            f"a razón de {horas_semanales} períodos lectivos semanales."
        )
    context["texto_introduccion"] = texto_intro

    texto_uds = config.get("texto_uds_modulo", "")
    if not texto_uds:
        texto_uds = f"El módulo de {modulo}, comprende las siguientes unidades formativas:"
    context["texto_uds_modulo"] = texto_uds

    # Lista de UDs en formato "UF{código}_{n} descripción" — un único bucle
    # dinámico (sustituye a los antiguos placeholders ud1_item..ud11_item,
    # que la plantilla nunca llegó a usar: los dos sitios donde aparecía
    # esta lista tenían el tag genérico sin numerar "{{ ud_item }}", que no
    # existía en el contexto y por tanto salía siempre en blanco).
    list_uf_items = [
        f"UF{codigo}_{i:02d} {resolve_ud_desc(ud, ud_desc_map)}"
        for i, ud in enumerate(df_ud, 1)
    ]
    context["list_uf_items"] = list_uf_items

    # ── SECCIÓN B: FEOE ────────────────────────────────────────────────
    texto_feoe = config.get("texto_feoe", "")
    if not texto_feoe:
        if df_ra:
            ra_feoe = [ra for ra in df_ra if ra.get('is_dual')]
            if ra_feoe:
                ids = ", ".join(_ra_id_full(ra) for ra in ra_feoe)
                texto_feoe = f"Los RA susceptibles de ser adquiridos en FEOE son: {ids}."
            else:
                texto_feoe = "No hay ningún RA dualizado."
        else:
            texto_feoe = "No hay ningún RA dualizado."
    context["texto_feoe"] = texto_feoe

    # ── SECCIÓN C: SECUENCIACIÓN — Tabla RA×UD (construida con python-docx
    # en generate(), no vía Jinja: así admite tantas filas de UD y columnas
    # de RA como el módulo tenga realmente, en vez del máximo fijo de la
    # plantilla original (11 UD × 7 RA) ─────────────────────────────────
    ra_ids = [_ra_id_full(ra) for ra in df_ra]
    context["tabla_secuenciacion"] = {
        "ra_ids": ra_ids,
        "ra_pcts": [str(int(ra.get('peso_ra', 0) or 0)) for ra in df_ra],
        "filas_ud": [
            {
                "num": str(i),
                "ev": f"{ev_map.get(str(ud.get('id_ud', '')))}ª" if ev_map.get(str(ud.get('id_ud', ''))) else "",
                "nombre": resolve_ud_desc(ud, ud_desc_map),
                "horas": str(ud.get('horas_ud', '')),
                "ra_vals": [
                    (str(int(ud.get(rid, 0))) if ud.get(rid, 0) else "")
                    for rid in ra_ids
                ],
            }
            for i, ud in enumerate(df_ud, 1)
        ],
    }

    # ── SECCIÓN C1: CONTENIDOS — reutiliza el mismo bucle dinámico
    # {%p for ud in list_uds %} que ya trae la plantilla, con el formato
    # "UD n. descripción" que espera esta sección ────────────────────────
    context["list_uds"] = [f"UD {i}. {resolve_ud_desc(ud, ud_desc_map)}" for i, ud in enumerate(df_ud, 1)]

    # ── SECCIÓN C2: CRITERIOS DE EVALUACIÓN — RA + sus CE, número variable
    # de RA y de CE por RA. Se aplana en una única lista de líneas (cada una
    # su propio párrafo vía {%p for %}) en vez de anidar bucles Jinja. ─────
    ce_by_ra: dict = {}
    for ce in df_ce:
        ce_by_ra.setdefault(str(ce.get("id_ra", "")), []).append(ce)

    # Catálogo oficial agrupado por RA (curriculo_data, resuelto en
    # routers/pdf.py vía fetch_curriculo_from_db) — hace falta como
    # respaldo POSICIONAL, no solo por id: el df_ce local suele numerar los
    # CE con letras ("CE1.a", "CE1.b"...) y el catálogo con números
    # ("CE1.1", "CE1.2"...), así que aunque sean los mismos criterios el id
    # normalizado nunca coincide entre uno y otro.
    catalog_ce_by_ra: dict = {}
    for ra_cat in (data.get("curriculo_data") or {}).get("ra") or []:
        rid_cat = str(ra_cat.get("id", "")).strip().rstrip(".")
        catalog_ce_by_ra[rid_cat] = ra_cat.get("ce") or []

    list_c2 = []
    for ra in df_ra:
        rid = str(ra.get("id_ra", ""))
        list_c2.append(f"{_ra_id_full(ra)}. {resolve_ra_desc(ra, ra_desc_map)}")
        list_c2.append("Criterios de evaluación:")

        ces_locales = ce_by_ra.get(rid, [])
        ces_catalogo = catalog_ce_by_ra.get(rid, [])
        mismo_recuento = len(ces_locales) == len(ces_catalogo) and len(ces_catalogo) > 0

        for i, ce in enumerate(ces_locales):
            desc = resolve_ce_desc(ce, ce_desc_map)
            if not desc and mismo_recuento:
                desc = ces_catalogo[i].get("descripcion", "")
            if desc:
                peso = ce.get("peso_ce")
                prefijo = f"({int(peso)}%) " if peso else ""
                list_c2.append(f"{prefijo}{desc}")
    context["list_c2"] = list_c2

    # ── SECCIÓN C3: CRITERIOS DE CALIFICACIÓN ──────────────────────────
    texto_calif = config.get("texto_criterios_calificacion", "")
    if not texto_calif:
        texto_calif = "30%. Conceptuales. 30%. Procedimentales. 40%. Actitudinales."
    context["texto_criterios_calificacion"] = texto_calif

    # Misma tabla "Instrumento | 1er/2º/3er Tri." que PD- (helpers_pd_tablas)
    instrumentos_pct = data.get("instrumentos_pct_trimestre") or []
    if not instrumentos_pct:
        instrumentos_pct = [
            {"nombre": "Exámenes teóricos", "pct_1t": 30, "pct_2t": 20, "pct_3t": 10},
            {"nombre": "Exámenes prácticos", "pct_1t": 20, "pct_2t": 20, "pct_3t": 10},
            {"nombre": "Exposición y defensa proyecto", "pct_1t": 10, "pct_2t": 20, "pct_3t": 30},
            {"nombre": "Informes de ejercicios", "pct_1t": 20, "pct_2t": 30, "pct_3t": 40},
            {"nombre": "Cuaderno de tareas", "pct_1t": 20, "pct_2t": 10, "pct_3t": 10},
        ]
    context["list_instrumentos"] = [
        {
            "nombre": row.get("nombre", ""),
            "pct_1t": f"{row.get('pct_1t', 0)}%",
            "pct_2t": f"{row.get('pct_2t', 0)}%",
            "pct_3t": f"{row.get('pct_3t', 0)}%",
        }
        for row in instrumentos_pct
    ]
    context["pond_1t"] = info_mod.get("pond_1t", 30)
    context["pond_2t"] = info_mod.get("pond_2t", 30)
    context["pond_3t"] = info_mod.get("pond_3t", 40)

    # ── SECCIÓN H: RESULTADOS DE APRENDIZAJE — número variable de RA ────
    context["n_ra"] = len(df_ra)
    context["list_ras_h"] = [f"RA {i}. {resolve_ra_desc(ra, ra_desc_map)}" for i, ra in enumerate(df_ra, 1)]

    # ── SECCIÓN J: HERRAMIENTAS — resuelve recursos_espacios (catálogo
    # codificado en helpers_catalogo.RECURSOS_LABELS + añadidos libres del
    # profesorado) ────────────────────────────────────────────────────────
    recursos = data.get("recursos_espacios") or []
    context["list_herramientas"] = resolve_recursos(recursos) if recursos else DEFAULT_HERRAMIENTAS

    # ── SECCIÓN N: PLAN DE CONTINGENCIA — coordinación con otros módulos,
    # mismo campo textos_pd_metodologia_labor_coordinada que se podrá
    # reutilizar en PD+ (ver docs/pd-plus-pendientes.md) ─────────────────
    context["textos_pd_metodologia_labor_coordinada"] = (
        data.get("textos_pd_metodologia_labor_coordinada") or DEFAULT_TEXTO_COORDINACION
    )

    return context


def generate(data: dict, out_docx: str, out_pdf: str = None):
    """
    Genera la PD= (Suficiente/BOA) usando la plantilla DOCX.
    """
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(
            f"No se encontró la plantilla en: {TEMPLATE_PATH}. "
            f"Ejecuta 'python scripts/preparar_plantilla_pd_suficiente.py' para generarla."
        )

    tpl = DocxTemplate(TEMPLATE_PATH)
    context = _build_context(data)
    tpl.render(context)

    doc = tpl.docx
    _insertar_tabla_secuenciacion(doc, context["tabla_secuenciacion"])
    for p in doc.paragraphs:
        if "[[TABLA_INSTRUMENTOS_C3]]" in p.text:
            p.text = p.text.replace("[[TABLA_INSTRUMENTOS_C3]]", "")
            insertar_tabla_instrumentos(
                doc, p, context["list_instrumentos"],
                context["pond_1t"], context["pond_2t"], context["pond_3t"],
            )
            break

    tpl.save(out_docx)


def _insertar_tabla_secuenciacion(doc, tabla_data: dict):
    """
    Construye con python-docx la tabla de la sección C (Secuenciación),
    en el marcador [[TABLA_SECUENCIACION]]: N | Ev. | Unidades didácticas |
    H. | RA1..RAn (tantas columnas de RA como tenga el módulo) y una fila
    de "% por RA" además de una fila por cada UD.
    """
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from helpers_pd_tablas import aplicar_bordes_rejilla

    ra_ids = tabla_data["ra_ids"]
    ra_pcts = tabla_data["ra_pcts"]
    filas_ud = tabla_data["filas_ud"]
    n_cols = 4 + len(ra_ids)

    for p in doc.paragraphs:
        if "[[TABLA_SECUENCIACION]]" not in p.text:
            continue
        p.text = p.text.replace("[[TABLA_SECUENCIACION]]", "")

        table = doc.add_table(rows=2 + len(filas_ud), cols=n_cols)
        aplicar_bordes_rejilla(table)

        # Anchos: 18cm totales repartidos entre las 4 columnas fijas (más
        # anchas) y las columnas de RA (más estrechas), igual de estrecho
        # cuantas más columnas de RA haya para no desbordar la página.
        fixed_widths = [Cm(1.0), Cm(1.2), Cm(7.0), Cm(1.2)]
        resto = 18.0 - sum(w.cm for w in fixed_widths)
        ra_w = Cm(max(resto / len(ra_ids), 0.8)) if ra_ids else Cm(0)
        col_widths = fixed_widths + [ra_w] * len(ra_ids)
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w

        def _fmt_cell(cell, text, bold=False, center=True):
            cell.text = str(text)
            for para in cell.paragraphs:
                if center:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in para.runs:
                    r.bold = bold
                    r.font.name = 'Arial'
                    r.font.size = Pt(8)
                    r.font.color.rgb = RGBColor(0, 0, 0)

        # Fila 0: "% por RA" en las columnas de RA
        pct_cells = table.rows[0].cells
        _fmt_cell(pct_cells[2], "% por RA", bold=True, center=False)
        for j, pct in enumerate(ra_pcts):
            _fmt_cell(pct_cells[4 + j], f"{pct}%", bold=True)

        # Fila 1: cabeceras de columna
        hdr_cells = table.rows[1].cells
        for i, label in enumerate(["N", "Ev.", "Unidades didácticas", "H."]):
            _fmt_cell(hdr_cells[i], label, bold=True, center=(i != 2))
        for j, rid in enumerate(ra_ids):
            _fmt_cell(hdr_cells[4 + j], rid, bold=True)

        # Filas de datos, una por UD
        for i, fila in enumerate(filas_ud):
            row_cells = table.rows[2 + i].cells
            _fmt_cell(row_cells[0], fila["num"])
            _fmt_cell(row_cells[1], fila["ev"])
            _fmt_cell(row_cells[2], fila["nombre"], center=False)
            _fmt_cell(row_cells[3], fila["horas"])
            for j, val in enumerate(fila["ra_vals"]):
                _fmt_cell(row_cells[4 + j], val)

        p._p.addnext(table._tbl)
        p._p.getparent().remove(p._p)
        break
